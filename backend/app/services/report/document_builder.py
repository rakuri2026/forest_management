"""
Document builder - Generate .docx report from AI-generated sections
"""
import os
import re
import tempfile
from typing import Dict, Any, Optional, List
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.utils.number_format import format_devanagari


def build_cover_page(doc: Document, metadata: Dict) -> None:
    """Build the cover page of the report"""
    # Add spacing at top
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("सामुदायिक वन कार्य योजना")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("COMMUNITY FOREST OPERATIONAL PLAN")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # Spacing
    doc.add_paragraph()

    # Details table
    details = [
        ("क्रम संख्या (Serial No.)", metadata.get('serial_number', '...............')),
        ("सामुदायिक वनको कोड (CF Code)", metadata.get('cf_code', '...............')),
        ("प्रदेश/डिभिजन/सब डिभिजन/पालिका", f"{metadata.get('province', '.....')} / {metadata.get('division', '..........')} / {metadata.get('sub_division', '..............')} / {metadata.get('municipality', '.........')}"),
        ("सामुदायिक वनको नाम", metadata.get('forest_name', '')),
        ("उपभोक्ता समूहको नाम", metadata.get('group_name', '')),
        ("ठेगाना", metadata.get('address', '')),
    ]

    table = doc.add_table(rows=len(details), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(details):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)

        cell_label.text = label
        cell_value.text = str(value)

        for cell in [cell_label, cell_value]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    if cell == cell_label:
                        run.font.bold = True

    # Period
    doc.add_paragraph()
    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period.add_run(f"आ.व. {metadata.get('fy_start', '२०../..')} देखि आ.व. {metadata.get('fy_end', '२०../..')} सम्म")
    run.font.size = Pt(14)
    run.font.bold = True

    period_en = doc.add_paragraph()
    period_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period_en.add_run(f"FY {metadata.get('fy_start', '20../..')} TO FY {metadata.get('fy_end', '20../..')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # National database code
    doc.add_paragraph()
    code = doc.add_paragraph()
    code.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = code.add_run(f"CF National Database Code: {metadata.get('cf_national_code', '...............')}")
    run.font.size = Pt(11)

    doc.add_page_break()


def add_section_heading(doc: Document, section_num: str, title_ne: str, title_en: str = "") -> None:
    """Add section heading"""
    heading = doc.add_heading(f'{section_num}. {title_ne}', level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 80, 0)

    if title_en:
        sub = doc.add_paragraph()
        run = sub.add_run(title_en)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(120, 120, 120)
        sub.paragraph_format.space_after = Pt(12)


def add_subsection_heading(doc: Document, section_num: str, subsection: str, title_ne: str) -> None:
    """Add subsection heading"""
    heading = doc.add_heading(f'{section_num}({subsection}) {title_ne}', level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 100, 0)
    heading.paragraph_format.space_before = Pt(18)


_NUM_PATTERN = re.compile(r"\b\d+(?:\.\d+)?\b")


def _replace_numbers(text: str) -> str:
    """Replace all numeric values in text with Devanagari digits."""
    def _replace(m):
        return format_devanagari(m.group(0), 2)
    return _NUM_PATTERN.sub(_replace, text)


def add_text_content(doc: Document, text: str, language: str = "NP") -> None:
    """Add text content with proper formatting"""
    # Clean up the text
    text = text.strip()

    # Split by paragraphs
    paragraphs = text.split('\n')

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        if language == "NP":
            para_text = _replace_numbers(para_text)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15

        run = p.add_run(para_text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'


def add_table_from_dict(doc: Document, headers: List[str], rows: List[List], caption: str = "") -> None:
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)

        # Set header background color
        shading = cell._tc.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '006400',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elem)

    # Data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = format_devanagari(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True


def add_image(doc: Document, image_path_or_data: str, caption: str = "", width: float = 5.5) -> None:
    """Add an image to the document"""
    if not image_path_or_data:
        return

    try:
        if image_path_or_data.startswith('data:'):
            # Base64 encoded image (SVG or PNG)
            from app.utils.svg_to_png import add_svg_picture
            add_svg_picture(doc, image_path_or_data, width_inches=width)
        else:
            # File path
            with open(image_path_or_data, "rb") as _f:
                img_bytes = _f.read()
            from app.utils.svg_to_png import add_svg_picture
            add_svg_picture(doc, img_bytes, width_inches=width)

        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()  # Spacing after image

    except Exception as e:
        # If image fails, add placeholder
        p = doc.add_paragraph()
        run = p.add_run(f"[Image: {caption}]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(200, 0, 0)


def build_report_document(
    metadata: Dict,
    sections: Dict[str, Dict],
    include_images: bool = True,
) -> BytesIO:
    """Build the complete report document

    Args:
        metadata: Report metadata (cover page info)
        sections: Dict of section_number -> {title_ne, title_en, subsections: {key: {title_ne, content, images}}}
        include_images: Whether to include map/chart images

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cover page
    build_cover_page(doc, metadata)

    # Table of Contents placeholder
    toc_heading = doc.add_heading('विषय सूची (Table of Contents)', level=1)
    for run in toc_heading.runs:
        run.font.color.rgb = RGBColor(0, 80, 0)

    p = doc.add_paragraph()
    run = p.add_run("[Table of Contents will be generated automatically in MS Word. Use Ctrl+A then F9 to update.]")
    run.font.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(10)
    doc.add_page_break()

    # Add each section
    def sort_key(x):
        try:
            parts = x.split('.')
            if parts[0].isdigit():
                return (int(parts[0]), int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0)
            return (99, 0)
        except (ValueError, IndexError):
            return (99, 0)

    doc_language = (metadata.get("language") or metadata.get("plan_language") or "NP").upper()

    for section_num in sorted(sections.keys(), key=sort_key):
        section_data = sections[section_num]
        title_ne = section_data.get('title_ne', '')
        title_en = section_data.get('title_en', '')

        # Check if this section has subsections
        if 'subsections' in section_data:
            add_section_heading(doc, section_num, title_ne, title_en)

            for sub_key, sub_data in section_data['subsections'].items():
                add_subsection_heading(doc, section_num, sub_key, sub_data.get('title_ne', ''))

                content = sub_data.get('content', '')
                if content:
                    add_text_content(doc, content, language=doc_language)

                if include_images:
                    images = sub_data.get('images', [])
                    for img in images:
                        add_image(doc, img.get('data', ''), img.get('caption', ''))
        else:
            add_section_heading(doc, section_num, title_ne, title_en)

            content = section_data.get('content', '')
            if content:
                add_text_content(doc, content, language=doc_language)

            if include_images:
                images = section_data.get('images', [])
                for img in images:
                    add_image(doc, img.get('data', ''), img.get('caption', ''))

        doc.add_page_break()

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
