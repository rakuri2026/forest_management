"""
DOCX builder for Operational Plan export
Walks the resolved tree and builds a .docx document with headings, text, charts, and tables.
"""
import logging
import os
import time
from typing import Dict, Any, Optional, List
from io import BytesIO
from uuid import UUID
from unicodedata import normalize as _norm
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy.orm import joinedload

from app.models.op_table import OPTableData
from app.models.forest_block import ForestBlock
from app.models.calculation import Calculation
from app.models.field_inventory import FieldInventoryMeasurement, FieldInventorySamplePlot, FieldInventoryCalculation
from app.services.operational_plan.tree_models import TreeNode
from app.services.operational_plan.variable_resolver import VariableResolver
from app.services.operational_plan.variable_registry import get_variable
from app.utils.number_format import format_devanagari

# ── Chart SVG cache ──
_CHART_CACHE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "uploads", "charts_cache")

def _chart_cache_path(calculation_id: UUID, chart_key: str) -> str:
    sub = os.path.join(_CHART_CACHE_DIR, str(calculation_id))
    os.makedirs(sub, exist_ok=True)
    return os.path.join(sub, f"{chart_key}.svg")

def _chart_cache_get(calculation_id: UUID, chart_key: str) -> Optional[BytesIO]:
    path = _chart_cache_path(calculation_id, chart_key)
    if not os.path.exists(path):
        old = path.replace(".svg", ".png")
        if os.path.exists(old):
            path = old
    if os.path.exists(path):
        try:
            with open(path, "rb") as f:
                return BytesIO(f.read())
        except Exception:
            pass
    return None

def _chart_cache_set(calculation_id: UUID, chart_key: str, buf: BytesIO):
    path = _chart_cache_path(calculation_id, chart_key)
    try:
        with open(path, "wb") as f:
            f.write(buf.getvalue())
    except Exception:
        pass

def clear_chart_cache(calculation_id: UUID = None, chart_key: str = None):
    """Clear chart SVG cache. If calculation_id is None, clears all charts.
       If chart_key is also given, clears only that specific chart."""
    import shutil
    if calculation_id:
        sub = os.path.join(_CHART_CACHE_DIR, str(calculation_id))
        if chart_key:
            path = os.path.join(sub, f"{chart_key}.svg")
            if os.path.exists(path):
                os.remove(path)
            old = os.path.join(sub, f"{chart_key}.png")
            if os.path.exists(old):
                os.remove(old)
        else:
            if os.path.exists(sub):
                shutil.rmtree(sub, ignore_errors=True)
    else:
        if os.path.exists(_CHART_CACHE_DIR):
            shutil.rmtree(_CHART_CACHE_DIR, ignore_errors=True)
            os.makedirs(_CHART_CACHE_DIR, exist_ok=True)

# ── Nepali font setup for matplotlib ──
_FONT_SETUP_DONE = False
_DEV_FONT_PATH = None
def _ensure_dev_font():
    global _FONT_SETUP_DONE, _DEV_FONT_PATH
    if _FONT_SETUP_DONE:
        return
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.font_manager as fm
    import matplotlib.pyplot as plt
    for name in ['Nirmala UI', 'Mangal', 'Arial Unicode MS', 'Noto Sans Devanagari']:
        try:
            fp = fm.findfont(name, fallback_to_default=False)
            if fp:
                fm.fontManager.addfont(fp)
                plt.rcParams['font.family'] = name
                _DEV_FONT_PATH = fp
                _FONT_SETUP_DONE = True
                return
        except Exception:
            continue
    _FONT_SETUP_DONE = True

def _dev_fontprop(size: int = 14):
    if _DEV_FONT_PATH:
        import matplotlib.font_manager as fm
        return fm.FontProperties(fname=_DEV_FONT_PATH, size=size)
    return fm.FontProperties(size=size)

# ── Chart color maps (consistent with frontend) ──
_HEALTH_COLORS_MAP = {
    "excellent": "#228B22", "healthy": "#90EE90", "moderate": "#FFD700",
    "poor": "#FF8C00", "stressed": "#DC143C",
}
_ASPECT_COLORS_MAP = {
    "N": "#1A5490", "NE": "#3498DB", "E": "#1ABC9C", "SE": "#F1C40F",
    "S": "#E74C3C", "SW": "#E67E22", "W": "#F39C12", "NW": "#9B59B6",
    "Flat": "#CCCCCC",
}
_NASA_FOREST_QUALITY_COLORS = {
    "Primary Forest": "#00FF00", "Young Secondary Forest": "#FF0000",
    "Old Secondary Forest": "#6666FF",
}
_SOIL_BAR_COLORS = ["#8B4513", "#A0522D", "#CD853F", "#D2691E", "#DEB887", "#D2B48C"]

# Chart/map generators loaded lazily inside functions to avoid slow matplotlib import at module level


_COVER_GREEN = RGBColor(0, 100, 0)
_SUBTITLE_GRAY = RGBColor(80, 80, 80)

import docx.document as _docx_doc

# Patch Document.add_table to auto-repeat header row on every page
_orig_add_table = _docx_doc.Document.add_table
def _patched_add_table(self, rows, cols, *args, **kwargs):
    tbl = _orig_add_table(self, rows, cols, *args, **kwargs)
    if rows > 0:
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:tblHeader"))
    return tbl
_docx_doc.Document.add_table = _patched_add_table


def _set_cell_shading(cell, color_hex: str):
    shading = cell._tc.get_or_add_tcPr()
    elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(elem)


def _set_page_border(section):
    """Add double-line green page border to a section."""
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), "006400")
        pgBorders.append(border)
    sectPr.append(pgBorders)


def _add_cover_page(doc: Document, plan: Dict[str, Any], metadata: Dict[str, Any]):
    section = doc.sections[0]
    _set_page_border(section)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    user_inputs = metadata.get("user_inputs", {})
    forest_name = plan.get("forest_name", user_inputs.get("forest_name", "..............."))
    sn_number = user_inputs.get("sn_number", "...............")
    division = user_inputs.get("division", "...........")
    sub_division = user_inputs.get("sub_division", "...........")
    forest_municipality = user_inputs.get("forest_municipality", "...........")
    forest_ward = user_inputs.get("forest_ward", "...........")
    op_prep_year = user_inputs.get("op_preparation_year", "...........")
    op_start_fy = user_inputs.get("op_start_fy", "२०../..")
    op_end_fy = user_inputs.get("op_end_fy", "२०../..")
    cf_reg_no = user_inputs.get("cf_registration_number", "")
    cf_handover = user_inputs.get("cf_handover_date", "")
    district = user_inputs.get("district", "")

    np = _fix

    # -- top spacer --
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)

    # -- registration number (right) --
    if cf_reg_no:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"सामुदायिक वन द.नं. : {np(cf_reg_no)}")
        r.font.size = Pt(11)
        r.font.bold = True

    # -- main title --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("सामुदायिक वन व्यवस्थापन कार्ययोजना")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = _COVER_GREEN

    # -- year --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(np(str(op_prep_year)))
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = _COVER_GREEN

    # -- divider --
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "006400")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # -- details section --
    details = [
        ("क्रम संख्या :", np(str(sn_number))),
        ("डिभिजन / सव डिभिजन / स्थानिय तह / वडा :",
         f"{np(division)} / {np(sub_division)} / {np(forest_municipality)} / {np(forest_ward)}"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(20)
        r = p.add_run(label + " ")
        r.font.size = Pt(11)
        r.font.bold = True
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    # -- group name --
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(np(f"श्री {forest_name} सामुदायिक वन उपभोक्ता समूह"))
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = _COVER_GREEN

    # -- address --
    mun_label = user_inputs.get("municipality_type", "")
    mun_display = f"{forest_municipality} {mun_label}" if mun_label else forest_municipality
    addr_parts = [f"{mun_display} वडा नं. {forest_ward}"]
    if district:
        addr_parts.append(district)
    address = ", ".join(addr_parts)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(np(address))
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(60, 60, 60)

    # -- FY period --
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(np(f"आ.व. {op_start_fy} देखि आ.व. {op_end_fy} सम्म"))
    r.font.size = Pt(11)
    r.font.bold = True

    # -- English section (bordered) --
    en_table = doc.add_table(rows=1, cols=1)
    en_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    en_cell = en_table.cell(0, 0)

    # style the cell with double border
    tcPr = en_cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "006400")
        tcBorders.append(border)
    tcPr.append(tcBorders)

    # set cell paragraph alignment center
    for para in en_cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    en_cell.paragraphs[0].paragraph_format.space_before = Pt(6)
    en_cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    r = en_cell.paragraphs[0].add_run("COMMUNITY FORESTRY OPERATIONAL PLAN")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(60, 60, 60)

    p2 = en_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    r = p2.add_run(f"FY {op_start_fy} TO {op_end_fy}")
    r.font.size = Pt(10)
    r.font.color.rgb = _SUBTITLE_GRAY

    p3 = en_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(1)
    p3.paragraph_format.space_after = Pt(1)
    r = p3.add_run(f"Serial No.: {sn_number}")
    r.font.size = Pt(10)
    r.font.bold = True

    p4 = en_cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(1)
    r = p4.add_run(f"District/Sub Division/RM/Ward: {division} / {sub_division} / {forest_municipality} / {forest_ward}")
    r.font.size = Pt(10)

    p5 = en_cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(2)
    p5.paragraph_format.space_after = Pt(4)
    r = p5.add_run(f"Shree {forest_name} Community Forest User Group")
    r.font.size = Pt(10)
    r.font.bold = True

    # -- spacer to push bottom --
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # -- handover date --
    if cf_handover:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(np(f"वन हस्तान्तरण मिति : {cf_handover}"))
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(60, 60, 60)

    # -- footer title --
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(np(f"वन व्यवस्थापन कार्ययोजना {op_prep_year}"))
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = _COVER_GREEN

    doc.add_page_break()


def _add_toc_field(doc: Document):
    heading = doc.add_heading("विषय सूची (Table of Contents)", level=1)
    for r in heading.runs:
        r.font.color.rgb = _COVER_GREEN

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    from docx.oxml import OxmlElement

    def _make_run_with_fldchar(para, fldchar_type):
        run = OxmlElement("w:r")
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), fldchar_type)
        run.append(fld)
        para._p.append(run)

    def _make_run_with_instr(para, text):
        run = OxmlElement("w:r")
        instr = OxmlElement("w:instrText")
        instr.text = text
        run.append(instr)
        para._p.append(run)

    def _make_run_with_text(para, text, italic=False, color=None, size=None):
        run_elem = OxmlElement("w:r")
        t_elem = OxmlElement("w:t")
        t_elem.text = text
        run_elem.append(t_elem)
        if italic or color or size:
            rpr = OxmlElement("w:rPr")
            if italic:
                rpr.append(OxmlElement("w:i"))
            if color:
                c = OxmlElement("w:color")
                c.set(qn("w:val"), color)
                rpr.append(c)
            if size:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(size))
                rpr.append(sz)
            run_elem.insert(0, rpr)
        para._p.append(run_elem)

    _make_run_with_fldchar(p, "begin")
    _make_run_with_instr(p, ' TOC \\o "1-3" \\h \\z \\u ')
    _make_run_with_fldchar(p, "separate")
    _make_run_with_text(p, "[Right-click here and select 'Update Field' to generate Table of Contents]",
                        italic=True, color="969696", size=20)
    _make_run_with_fldchar(p, "end")

    doc.add_page_break()


def _add_heading(doc: Document, node: TreeNode):
    num = f"{node.number}. " if node.number else ""
    text = f"{num}{_norm('NFC', node.title_ne)}"

    if node.type in ("section", "appendix"):
        heading = doc.add_heading(text, level=1)
    elif node.type == "subsection":
        heading = doc.add_heading(text, level=2)
    elif node.type == "preamble":
        heading = doc.add_heading(text, level=1)
    else:
        heading = doc.add_heading(text, level=1)

    for r in heading.runs:
        r.font.color.rgb = _COVER_GREEN

    if node.title_en and node.title_en != node.title_ne:
        sub = doc.add_paragraph()
        run = sub.add_run(_fix( node.title_en))
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(120, 120, 120)
        sub.paragraph_format.space_after = Pt(12)


import re

_VARIABLE_PATTERN = re.compile(r"\{\{(\w+(?::\w+)*)\}\}")

# Devanagari text cleanup: common decomposed/wrong sequences → proper composed forms
_DEVANAGARI_FIXES = [
    ("\u093E\u0948", "\u094C"),  # ा+ै → ौ
    ("\u0947\u0947", "\u0947"),  # े+े → े (duplicate e)
    ("\u0948\u0948", "\u0948"),  # ै+ै → ै (duplicate ai)
]

_ARABIC_TO_DEV = str.maketrans("0123456789", "०१२३४५६७८९")

def _fix(text: str) -> str:
    """Normalize NFC + fix common Devanagari sequence errors + convert Arabic digits to Devanagari."""
    if text is None:
        return ""
    text = _norm("NFC", text)
    for wrong, correct in _DEVANAGARI_FIXES:
        text = text.replace(wrong, correct)
    text = text.translate(_ARABIC_TO_DEV)
    return text


def _set_header_repeat(tbl, row_index=0):
    """Mark a table row to repeat as header on each page when table spans multiple pages."""
    trPr = tbl.rows[row_index]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))

def _fmt_value(value, var_name=""):
    """Format a resolved variable value with Devanagari digits and correct precision."""
    if value is None:
        return "-"
    if isinstance(value, str) and any(c in "०१२३४५६७८९" for c in value):
        return _fix(value)
    val = get_variable(var_name)
    precision = val.precision if val else 2
    return _fix( format_devanagari(value, precision))


def _resolve_var_text(text: str, raw_data: dict) -> str:
    """Resolve {{variable}} patterns in a text string using raw_data."""
    def _replacer(m):
        var_name = m.group(1)
        if var_name.startswith("chart:") or var_name.startswith("map:") or var_name.startswith("table:") or var_name.endswith(":full"):
            return m.group(0)
        var_val = _resolve_var_from_raw(var_name, raw_data)
        if var_val is None:
            return ""
        if isinstance(var_val, (dict, list)):
            return ""
        return _fmt_value(var_val, var_name)
    return _fix( re.sub(r"\{\{(\w+(?::\w+)*)\}\}", _replacer, text))

# Alias mapping to look up unresolved list/dict variables from raw_data
_VAR_LOOKUP = {
    # Boundary features
    "boundary_features_north": ("boundary", "features.north"),
    "boundary_features_east": ("boundary", "features.east"),
    "boundary_features_south": ("boundary", "features.south"),
    "boundary_features_west": ("boundary", "features.west"),
    "extent_n": ("boundary", "whole_forest_extent.N"),
    "extent_s": ("boundary", "whole_forest_extent.S"),
    "extent_e": ("boundary", "whole_forest_extent.E"),
    "extent_w": ("boundary", "whole_forest_extent.W"),

    # User group
    "ug_land_cover_classes": ("user_group", "land_cover_classes"),

    "uc_members": ("committees", "user_committee.members"),
    "uc_gender_distribution": ("committees", "user_committee.gender_distribution"),
    "uc_position_distribution": ("committees", "user_committee.position_distribution"),
    "uc_caste_distribution": ("committees", "user_committee.caste_distribution"),
    "ac_members": ("committees", "advisory_committee.members"),
    "fc_members": ("committees", "financial_committee.members"),
    "species_list": ("species", "species_list"),
    "bio_vegetation": ("biodiversity", "vegetation"),
    "bio_animals": ("biodiversity", "animals"),
    "bio_available": ("biodiversity", "available"),
    "bio_total_species": ("biodiversity", "total_species"),
    "bio_vegetation_count": ("biodiversity", "vegetation_count"),
    "bio_animal_count": ("biodiversity", "animal_count"),
    "bio_protected_count": ("biodiversity", "protected_count"),
    "bio_invasive_count": ("biodiversity", "invasive_count"),
    "bio_iucn_cr": ("biodiversity", "iucn_breakdown.CR"),
    "bio_iucn_en": ("biodiversity", "iucn_breakdown.EN"),
    "bio_iucn_vu": ("biodiversity", "iucn_breakdown.VU"),
    "bio_sub_category_breakdown": ("biodiversity", "sub_category_breakdown"),
    "activities_list": ("activities", "activities"),
    "ya_year_summary": ("yearly_plan", "year_summary"),
    "ya_plan_matrix": ("yearly_plan", "plan_matrix"),
    "ya_program_budget": ("yearly_plan", "program_budget"),
    "ya_total_budget_by_year": ("yearly_plan", "total_budget_by_year"),
    "ya_total_ten_year_budget": ("yearly_plan", "total_ten_year_budget"),
    "ya_activity_plan_detail": ("yearly_plan", "activity_plan_detail"),
    "ug_buildings": ("user_group", "buildings"),
    "hh_prosperity_distribution": ("households", "prosperity_distribution"),
    "hh_caste_distribution": ("households", "caste_distribution"),
    "inventory_species_summary": ("inventory", "species_summary"),
    "inventory_block_summary": ("inventory", "block_summary"),
    "blocks_count": ("blocks", "total_blocks"),
    "species_by_role": ("species", "by_role"),
    "kabuliyatnama_date": ("user_inputs", "kabuliyatnama_date"),
    "kabuliyatnama_date_year": ("user_inputs", "kabuliyatnama_date"),
    "kabuliyatnama_date_month": ("user_inputs", "kabuliyatnama_date"),
    "kabuliyatnama_date_day": ("user_inputs", "kabuliyatnama_date"),
    "kabuliyatnama_date_sentence": ("user_inputs", "kabuliyatnama_date"),

    # Sampling
    "sampling_type": ("sampling", "designs.0.sampling_type"),
    "sampling_block_summary": ("sampling", "sampling_block_summary"),
    "sampling_point_locations": ("sampling", "sampling_point_locations"),
    "sampling_total_points": ("sampling", "designs.0.total_points"),
    "sampling_total_blocks": ("sampling", "designs.0.total_blocks"),
    "sampling_requested_intensity": ("sampling", "designs.0.requested_intensity_percent"),
    "sampling_actual_intensity": ("sampling", "designs.0.sampling_percentage"),
    "sampling_forest_area_ha": ("sampling", "designs.0.forest_area_hectares"),
    "sampling_plot_area_sqm": ("sampling", "designs.0.plot_area_sqm"),
    "sampling_total_sampled_area_ha": ("sampling", "designs.0.total_sampled_area_hectares"),
    "sampling_available": ("sampling", "available"),
    "sampling_plot_shape": ("sampling", "designs.0.plot_shape"),
    "sampling_plot_radius_m": ("sampling", "designs.0.plot_radius_meters"),
    "sampling_intensity_per_ha": ("sampling", "designs.0.intensity_per_hectare"),

    # Fieldbook
    "fieldbook_total_points": ("fieldbook", "total_points"),
    "fieldbook_vertex_count": ("fieldbook", "vertex_count"),
    "fieldbook_interpolated_count": ("fieldbook", "interpolated_count"),
    "fieldbook_perimeter_m": ("fieldbook", "perimeter_m"),
    "fieldbook_avg_elevation_m": ("fieldbook", "avg_elevation_m"),
    "fieldbook_min_elevation_m": ("fieldbook", "min_elevation_m"),
    "fieldbook_max_elevation_m": ("fieldbook", "max_elevation_m"),
    "fieldbook_points": ("fieldbook", "points"),
    "fieldbook_block_summary": ("fieldbook", "block_summary"),
    "fieldbook_narration": ("section_generators", "section:fieldbook_narration"),

    # Tree Mapping Analysis (sm_* variables)
    "sm_available":                    ("tree_mapping_analysis", "sm_available"),
    "sm_total_blocks_analyzed":        ("tree_mapping_analysis", "sm_total_blocks_analyzed"),
    "sm_total_trees_analyzed":         ("tree_mapping_analysis", "sm_total_trees_analyzed"),
    "sm_total_carbon_tc":             ("tree_mapping_analysis", "sm_total_carbon_tc"),
    "sm_total_co2_tco2":             ("tree_mapping_analysis", "sm_total_co2_tco2"),
    "sm_hierarchy_summary":            ("tree_mapping_analysis", "sm_hierarchy_summary"),
    "sm_species_by_hierarchy":         ("tree_mapping_analysis", "sm_species_by_hierarchy"),
    "sm_species_diversity":            ("tree_mapping_analysis", "sm_species_diversity"),
    "sm_dbh_by_hierarchy":             ("tree_mapping_analysis", "sm_dbh_by_hierarchy"),
    "sm_dbh_species_by_hierarchy":     ("tree_mapping_analysis", "sm_dbh_species_by_hierarchy"),
    "sm_stand_type_by_hierarchy":      ("tree_mapping_analysis", "sm_stand_type_by_hierarchy"),
    "sm_forest_structure_status":      ("tree_mapping_analysis", "sm_forest_structure_status"),
    "sm_carbon_by_hierarchy":          ("tree_mapping_analysis", "sm_carbon_by_hierarchy"),
    "sm_volume_by_hierarchy":          ("tree_mapping_analysis", "sm_volume_by_hierarchy"),
    "sm_top_species_by_volume":        ("tree_mapping_analysis", "sm_top_species_by_volume"),
    "sm_mother_tree_coverage":         ("tree_mapping_analysis", "sm_mother_tree_coverage"),
    "sm_mother_tree_by_hierarchy":     ("tree_mapping_analysis", "sm_mother_tree_by_hierarchy"),
    "sm_mother_tree_by_species":       ("tree_mapping_analysis", "sm_mother_tree_by_species"),
    "sm_felling_tree_by_species":      ("tree_mapping_analysis", "sm_felling_tree_by_species"),
    "sm_mother_felling_summary":       ("tree_mapping_analysis", "sm_mother_felling_summary"),
    "sm_hierarchy_remark_breakdown":   ("tree_mapping_analysis", "sm_hierarchy_remark_breakdown"),
    "sm_species_hier_remark":          ("tree_mapping_analysis", "sm_species_hier_remark"),
    "sm_dbh_hier_remark":              ("tree_mapping_analysis", "sm_dbh_hier_remark"),
    "sm_felling_dbh_analysis":         ("tree_mapping_analysis", "sm_felling_dbh_analysis"),
    "sm_felling_species_analysis":     ("tree_mapping_analysis", "sm_felling_species_analysis"),
    "sm_felling_totals":               ("tree_mapping_analysis", "sm_felling_totals"),

    # Tree Mapping Analysis Narrations
    "section:sm_hierarchy_narration":     ("section_generators", "section:sm_hierarchy_narration"),
    "section:sm_species_narration":       ("section_generators", "section:sm_species_narration"),
    "section:sm_dbh_narration":           ("section_generators", "section:sm_dbh_narration"),
    "section:sm_stand_type_narration":    ("section_generators", "section:sm_stand_type_narration"),
    "section:sm_carbon_narration":        ("section_generators", "section:sm_carbon_narration"),
    "section:sm_volume_narration":        ("section_generators", "section:sm_volume_narration"),
    "section:sm_mother_tree_narration":   ("section_generators", "section:sm_mother_tree_narration"),
    "section:sm_felling_narration":       ("section_generators", "section:sm_felling_narration"),

    # Tree Mapping Analysis Legend Variables
    "sm_mf_hierarchy_legend":     ("tree_mapping_analysis", "sm_mf_hierarchy_legend"),
    "sm_stand_type_legend":       ("tree_mapping_analysis", "sm_stand_type_legend"),
    "sm_carbon_legend":           ("tree_mapping_analysis", "sm_carbon_legend"),
    "sm_volume_legend":           ("tree_mapping_analysis", "sm_volume_legend"),
    "sm_mf_species_legend":       ("tree_mapping_analysis", "sm_mf_species_legend"),
    "sm_felling_species_legend":  ("tree_mapping_analysis", "sm_felling_species_legend"),
}

def _deep_get(data, path):
    if not path:
        return None
    current = data
    for key in path.split("."):
        if isinstance(current, dict):
            current = current.get(key)
        elif isinstance(current, (list, tuple)):
            try:
                idx = int(key)
                current = current[idx] if 0 <= idx < len(current) else None
            except (ValueError, IndexError):
                return None
        else:
            return None
        if current is None:
            return None
    return current

def _resolve_var_from_raw(var_name: str, raw_data: dict) -> any:
    lookup = _VAR_LOOKUP.get(var_name)
    if lookup:
        section, path = lookup
        section_data = raw_data.get(section, {})
        result = _deep_get(section_data, path)
        if result is not None:
            if var_name == "kabuliyatnama_date_year":
                parts = str(result).split("/")
                return int(parts[0]) if len(parts) == 3 else result
            if var_name == "kabuliyatnama_date_month":
                parts = str(result).split("/")
                return int(parts[1]) if len(parts) == 3 else result
            if var_name == "kabuliyatnama_date_day":
                parts = str(result).split("/")
                return int(parts[2]) if len(parts) == 3 else result
            if var_name == "kabuliyatnama_date_sentence":
                return _format_kabuliyatnama_sentence(str(result))
            return result

    user_inputs = raw_data.get("user_inputs", {})
    if var_name in user_inputs and user_inputs[var_name] is not None:
        return user_inputs[var_name]

    for section_key in ("basic_info", "raster_analysis", "boundary", "blocks",
                        "species", "inventory", "field_inventory", "fieldbook",
                        "sampling", "households", "committees", "biodiversity",
                        "activities", "user_group", "section_generators",
                        "compartment", "yearly_plan"):
        section_data = raw_data.get(section_key, {})
        if var_name in section_data:
            return section_data[var_name]
    return None


def _format_kabuliyatnama_sentence(raw: str) -> str:
    parts = raw.split("/")
    if len(parts) != 3:
        return raw
    try:
        from nepali_datetime import date as nepali_date
        y, m, d = int(parts[0]), int(parts[1]), int(parts[2])
        nd = nepali_date(y, m, d)
        MONTH_NAMES_NP = (None, "वैशाख", "जेष्ठ", "असार", "श्रावण", "भदौ", "आश्विन", "कार्तिक", "मंसिर", "पौष", "माघ", "फाल्गुण", "चैत्र")
        DAY_NAMES_NP = ("आइतबार", "सोमबार", "मङ्गलबार", "बुधबार", "बिहिबार", "शुक्रबार", "शनिबार")
        month_name = MONTH_NAMES_NP[m] if 1 <= m <= 12 else str(m)
        day_name = DAY_NAMES_NP[nd.weekday()]
        return f"ईति सम्वत {_fix(str(y))} साल {month_name} महिना {_fix(str(d))} गते रोज {day_name} शुभम् ।"
    except Exception:
        return raw

def _render_list_value_as_text(val: any, var_name: str = "") -> str:
    if not val:
        return ""
    if isinstance(val, list):
        if all(isinstance(v, str) for v in val):
            return ", ".join(val)
        if all(isinstance(v, dict) for v in val):
            if var_name == "uc_members":
                headers, rows = _build_uc_members_data(val)
                col_widths = [max(len(cell) for cell in [h] + [r[i] for r in rows]) for i, h in enumerate(headers)]
                lines = []
                sep = " | ".join(h.ljust(col_widths[i]) for i, h in enumerate(headers))
                lines.append(sep)
                lines.append("-+-".join("-" * col_widths[i] for i in range(len(headers))))
                for row in rows:
                    lines.append(" | ".join(cell.ljust(col_widths[i]) for i, cell in enumerate(row)))
                return "\n".join(lines)
            lines = []
            for i, item in enumerate(val, 1):
                parts = [_fmt_value(item.get(k, ""), var_name) for k in ("name", "position", "gender") if k in item]
                clean = " — ".join(p for p in parts if p.strip())
                if clean:
                    lines.append(f"{i}. {clean}")
            return "\n".join(lines)
        return "\n".join(f"• {v}" for v in val if v)
    if isinstance(val, dict):
        return "\n".join(f"{k}: {_fmt_value(v, var_name)}" for k, v in val.items() if v)
    return _fmt_value(val, var_name)

def _add_list_table(doc: Document, val: list, var_name: str = ""):
    if not val or not isinstance(val, list):
        return None
    if all(isinstance(v, str) for v in val):
        for v in val:
            p = doc.add_paragraph(f"• {v}")
            p.paragraph_format.space_after = Pt(2)
        return True
    if all(isinstance(v, dict) for v in val):
        vdef = get_variable(var_name)
        if vdef and vdef.label_ne:
            p = doc.add_paragraph()
            run = p.add_run(vdef.label_ne)
            run.font.size = Pt(11)
            run.font.bold = True
        raw_keys = list(val[0].keys())
        np_raw = NP_HEADERS_BIODIVERSITY.get(var_name)
        if np_raw and isinstance(np_raw[0], tuple):
            ordered = [k for k, _ in np_raw if k in raw_keys]
            extra = [k for k in raw_keys if k not in ordered]
            headers = ordered + extra
        else:
            headers = raw_keys
        num_cols = len(headers)
        num_rows = len(val) + 1
        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        _tbl_fill_data(tbl, headers, val, var_name)
        doc.add_paragraph()
        return True
    return None


def _tbl_fill_data(tbl, headers, rows_data, var_name=""):
    """Fill table data using direct lxml for speed (avoids slow tbl.cell() O(N²))."""
    tbl_elem = tbl._tbl
    tr_elems = tbl_elem.findall(qn('w:tr'))
    if not tr_elems:
        return
    # Header row
    header_tr = tr_elems[0]
    header_tcs = header_tr.findall(qn('w:tc'))
    np_raw = NP_HEADERS_BIODIVERSITY.get(var_name)
    np_map = {}
    if np_raw:
        if np_raw and isinstance(np_raw[0], tuple):
            np_map = {k: v for k, v in np_raw}
        else:
            np_map = {headers[i]: v for i, v in enumerate(np_raw) if i < len(headers)}
    for ci, h in enumerate(headers):
        if ci >= len(header_tcs):
            break
        tc = header_tcs[ci]
        display = np_map.get(h, h.replace("_", " ").title())
        _tc_set_text(tc, display)
        _tc_style(tc, bold=True, size=Pt(9), color=RGBColor(255, 255, 255), shading="006400")
    # Data rows
    for ri, row_val in enumerate(rows_data, 1):
        if ri >= len(tr_elems):
            break
        tr = tr_elems[ri]
        tcs = tr.findall(qn('w:tc'))
        for ci, h in enumerate(headers):
            if ci >= len(tcs):
                break
            tc = tcs[ci]
            _tc_set_text(tc, _fmt_value(row_val.get(h, ""), var_name))
            _tc_style(tc, bold=False, size=Pt(9))


def _tc_set_text(tc, text):
    """Set cell text via direct lxml (replaces all content)."""
    for child in list(tc):
        tc.remove(child)
    p_elem = OxmlElement('w:p')
    r_elem = OxmlElement('w:r')
    t_elem = OxmlElement('w:t')
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = text
    r_elem.append(t_elem)
    p_elem.append(r_elem)
    tc.append(p_elem)


def _tc_style(tc, bold=False, size=None, color=None, shading=None):
    """Style a table cell's runs via direct lxml."""
    for p_elem in tc.findall(qn('w:p')):
        for r_elem in p_elem.findall(qn('w:r')):
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r_elem.insert(0, rPr)
            if bold:
                b_elem = rPr.find(qn('w:b'))
                if b_elem is None:
                    b_elem = OxmlElement('w:b')
                    rPr.append(b_elem)
            if size is not None:
                sz_elem = rPr.find(qn('w:sz'))
                if sz_elem is None:
                    sz_elem = OxmlElement('w:sz')
                    rPr.append(sz_elem)
                sz_elem.set(qn('w:val'), str(int(size.pt * 2)) if hasattr(size, 'pt') else str(size))
            if color is not None:
                c_elem = rPr.find(qn('w:color'))
                if c_elem is None:
                    c_elem = OxmlElement('w:color')
                    rPr.append(c_elem)
                c_elem.set(qn('w:val'), str(color))
    if shading:
        _set_cell_shading_direct(tc, shading)


def _set_cell_shading_direct(tc, color_hex):
    """Apply cell shading via direct lxml (avoids python-docx overhead)."""
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)


NP_HEADERS_ACTIVITY_PLAN = [
    ("s_no", "सि.नं."),
    ("activity", "कृयाकलाप"),
    ("program", "कार्यक्रम"),
    ("unit", "इकाइ"),
    ("quantity_years", "वार्षिक\nपरिमाण"),
    ("budget_years", "वार्षिक\nबजेट"),
    ("total_budget", "जम्मा\nबजेट\n(हजार)"),
    ("location_type", "स्थान\nप्रकार"),
    ("location_details", "स्थान\nविवरण"),
    ("spatial_features", "स्थानिक\nफिचर"),
]

NP_COL_WIDTHS_CM = [
    1.0,   # सि.नं.
    5.0,   # कृयाकलाप
    2.2,   # कार्यक्रम
    1.0,   # इकाइ
    2.5,   # वार्षिक परिमाण
    2.5,   # वार्षिक बजेट
    1.5,   # जम्मा बजेट (हजार) — narrow, 2-3 digit value
    1.5,   # स्थान प्रकार
    3.2,   # स्थान विवरण
    2.2,   # स्थानिक फिचर
]

NP_HEADERS_HH_RECORDS = [
    ("घर_नं", "घर नं", 1.2, "number"),
    ("पुरुष_मुखिया", "पुरुष\nघरमुली", 2.8, "text"),
    ("महिला_मुखिया", "महिला\nघरमुली", 2.8, "text"),
    ("जात_वर्गीकरण", "जात\nवर्गीकरण", 2.0, "text"),
    ("पुरुष", "पुरुष", 0.8, "number"),
    ("महिला", "महिला", 0.8, "number"),
    ("ठेगाना", "ठेगाना", 2.0, "text"),
    ("गाई_गोरु", "गाई गोरु", 0.7, "number"),
    ("भैंसी", "भैंसी", 0.7, "number"),
    ("बाख्रा_भेडा", "बाख्रा भेडा", 0.7, "number"),
    ("जग्गा_क्षेत्रफल", "जग्गा", 0.9, "number"),
    ("जग्गा_एकाइ", "इकाइ", 1.1, "vertical"),
    ("घाँस_भारी", "घाँस भारी", 1.2, "number"),
    ("पोल", "पोल", 0.8, "number"),
    ("काठ_cft", "काठ Cft", 0.8, "number"),
    ("दाउरा_भारी", "दाउरा भारी", 1.2, "number"),
    ("ओछ्यान_भारी", "सोतर भारी", 1.2, "number"),
    ("समृद्धि_स्तर", "समृद्धि\nस्तर", 1.5, "text"),
    ("वन_पेशा", "वन\nपेशा", 1.2, "text"),
]

NP_HEADERS_BIODIVERSITY: Dict[str, List[tuple]] = {
    "table_20": [
        ("sn", "सि.नं."),
        ("name", "नाम"),
        ("scientific_name", "वैज्ञानिक नाम"),
        ("type", "प्रकार"),
        ("sub_category", "उप-प्रकार"),
        ("iucn_status", "IUCN स्थिति"),
        ("is_protected", "संरक्षित"),
        ("is_invasive", "मिचाहा"),
    ],
    "table_33": [
        ("iucn_code", "IUCN कोड"),
        ("nepali_label", "संरक्षण स्थिति"),
        ("count", "सङ्ख्या"),
    ],
    "table_34": [
        ("sn", "सि.नं."),
        ("name", "नाम"),
        ("scientific_name", "वैज्ञानिक नाम"),
        ("sub_category", "उप-प्रकार"),
        ("iucn_status", "IUCN स्थिति"),
    ],
    "table_35": [
        ("sn", "सि.नं."),
        ("name", "नाम"),
        ("scientific_name", "वैज्ञानिक नाम"),
        ("sub_category", "उप-प्रकार"),
        ("iucn_status", "IUCN स्थिति"),
    ],
    "table_36": [
        ("sn", "सि.नं."),
        ("name", "नाम"),
        ("scientific_name", "वैज्ञानिक नाम"),
        ("sub_category", "उप-प्रकार"),
        ("iucn_status", "IUCN स्थिति"),
        ("is_protected", "संरक्षित"),
        ("is_invasive", "मिचाहा"),
        ("primary_use", "प्रमुख प्रयोग"),
    ],
    "table_37": [
        ("sn", "सि.नं."),
        ("name", "नाम"),
        ("scientific_name", "वैज्ञानिक नाम"),
        ("sub_category", "उप-प्रकार"),
        ("iucn_status", "IUCN स्थिति"),
        ("is_protected", "संरक्षित"),
        ("is_invasive", "मिचाहा"),
        ("primary_use", "प्रमुख प्रयोग"),
    ],
    "demand_supply": [
        ("product", "उत्पादन किसिम"),
        ("demand", "माग"),
        ("cf_regular", "झिँजा दाउरा तथा घाँस संकलन"),
        ("cf_aah", "वार्षिक संकलन परिमाण"),
        ("private", "निजि क्षेत्रबाट उत्पादन"),
        ("total_supply", "जम्मा आपूर्ति"),
        ("deficit", "बचत तथा कमी"),
    ],

    # Tree Mapping Analysis tables
    "sm_hierarchy_summary": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("tree_count", "रूख गणना"),
        ("total_volume_m3", "जम्मा आयतन"),
        ("avg_dbh_cm", "औसत ब्यास से.मी."),
        ("avg_height_m", "औसत उचाइ मि."),
        ("area_ha", "क्षेत्रफल हे."),
        ("trees_per_ha", "रूख गणना प्रति हेक्टर"),
        ("volume_per_ha", "आयतन प्रति हेक्टर"),
        ("dominant_species", "मुख्य प्रजाती"),
    ],
    "sm_species_by_hierarchy": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("species", "प्रजाति"),
        ("local_name", "स्थानीय नाम"),
        ("tree_count", "रूख गणना"),
        ("hierarchy_percent", "क्षेत्रमा रहेको रूखको प्रतिशत"),
        ("timber_m3", "काठ घ.मी."),
        ("firewood_m3", "दाउरा घ.मी."),
        ("gross_volume_m3", "ग्रस आयतन घ.मी."),
        ("net_volume_m3", "नेट आयतन घ.मी."),
        ("volume_percent", "आयतन प्रतिशत"),
        ("avg_dbh_cm", "औसत ब्यास से.मी."),
    ],
    "sm_species_diversity": [
        ("block_name", "ब्लक"),
        ("species_richness", "प्रजाति समृद्धि"),
        ("shannon_index", "श्यानन सूचकांक"),
        ("evenness", "समानता"),
    ],
    "sm_dbh_by_hierarchy": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("dbh_class", "डिबियच क्लास"),
        ("tree_count", "रूख गणना"),
        ("timber_m3", "काठ घ.मी."),
        ("firewood_m3", "दाउरा घ.मी."),
        ("net_volume_m3", "नेट आयतन घ.मी."),
        ("gross_volume_m3", "ग्रस आयतन घ.मी."),
        ("hierarchy_percent", "क्षेत्रमा रहेको रूखको प्रतिशत"),
    ],
    "sm_dbh_species_by_hierarchy": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("dbh_class", "डिबियच क्लास"),
        ("species", "प्रजाति"),
        ("local_name", "स्थानीय नाम"),
        ("tree_count", "रूख गणना"),
        ("timber_m3", "काठ घ.मी."),
        ("firewood_m3", "दाउरा घ.मी."),
        ("gross_volume_m3", "ग्रस आयतन घ.मी."),
        ("net_volume_m3", "नेट आयतन घ.मी."),
        ("hierarchy_percent", "क्षेत्रमा रहेको रूखको प्रतिशत"),
        ("dbh_species_percent", "डिबियच-प्रजाति प्रतिशत"),
    ],
    "sm_stand_type_by_hierarchy": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("regeneration", "पुनरुत्पादन"),
        ("sapling", "लाथ्रा"),
        ("pole", "पोल"),
        ("tree", "रूख"),
        ("total", "जम्मा"),
        ("regeneration_percent", "पुनरुत्पादन प्रतिशत"),
        ("structure_status", "अवस्था"),
    ],
    "sm_carbon_by_hierarchy": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("gross_volume_m3", "ग्रस आयतन घ.मी."),
        ("wood_density", "भारित काठ घनत्व"),
        ("agb_t", "AGB (टन)"),
        ("bgb_t", "BGB (टन)"),
        ("biomass_t", "जैविक पदार्थ (टन)"),
        ("carbon_tc", "कार्बन (tC)"),
        ("co2_tco2", "CO₂e (tCO₂)"),
    ],
    "sm_volume_by_hierarchy": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("stem_volume_m3", "काण्डको आयतन घ.मी."),
        ("branch_volume_m3", "हाँगा विँगाको आयतन घ.मी."),
        ("total_volume_m3", "जम्मा आयतन घ.मी."),
        ("net_volume_m3", "नेट आयतन घ.मी."),
        ("firewood_m3", "दाउरा घ.मी."),
        ("firewood_chatta", "दाउरा चट्टा"),
    ],
    "sm_top_species_by_volume": [
        ("species", "प्रजाति"),
        ("local_name", "स्थानीय नाम"),
        ("total_volume_m3", "जम्मा आयतन घ.मी."),
        ("percent", "प्रतिशत"),
    ],
    "sm_mother_tree_by_hierarchy": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("grid_cells", "ग्रिड सेल"),
        ("mother_trees", "माँउ रूख"),
        ("felling_trees", "कटानी रूख"),
        ("coverage_ratio", "कभरेज अनुपात"),
    ],
    "sm_mother_tree_by_species": [
        ("species", "प्रजाति"),
        ("local_name", "स्थानीय नाम"),
        ("tree_count", "रूख गणना"),
        ("percent", "प्रतिशत"),
        ("timber_m3", "काठ घ.मी."),
        ("avg_dbh_cm", "औसत ब्यास से.मी."),
    ],
    "sm_felling_tree_by_species": [
        ("species", "प्रजाति"),
        ("local_name", "स्थानीय नाम"),
        ("tree_count", "रूख गणना"),
        ("percent", "प्रतिशत"),
        ("timber_m3", "काठ घ.मी."),
        ("avg_dbh_cm", "औसत ब्यास से.मी."),
    ],
    "sm_species_hier_remark": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("species", "प्रजाति"),
        ("local_name", "स्थानीय नाम"),
        ("remark", "टिप्पणी"),
        ("tree_count", "रूख गणना"),
        ("timber_m3", "काठ घ.मी."),
        ("gross_volume_m3", "ग्रस आयतन घ.मी."),
        ("avg_dbh_cm", "औसत ब्यास से.मी."),
    ],
    "sm_dbh_hier_remark": [
        ("compartment", "कम्पार्टमेन्ट"),
        ("sub_compartment", "सब कम्पार्टमेन्ट"),
        ("dbh_class", "डिबियच क्लास"),
        ("remark", "टिप्पणी"),
        ("tree_count", "रूख गणना"),
        ("timber_m3", "काठ घ.मी."),
        ("gross_volume_m3", "ग्रस आयतन घ.मी."),
    ],
    "sm_felling_dbh_analysis": [
        ("dbh_class", "डिबियच क्लास"),
        ("tree_count", "रूख गणना"),
        ("percent", "प्रतिशत"),
        ("timber_m3", "काठ घ.मी."),
        ("firewood_m3", "दाउरा घ.मी."),
        ("gross_volume_m3", "ग्रस आयतन घ.मी."),
        ("net_volume_m3", "नेट आयतन घ.मी."),
        ("fuelwood_m3", "इन्धन घ.मी."),
        ("fuelwood_chatta", "चट्टा घ.मी."),
        ("avg_dbh_cm", "औसत ब्यास से.मी."),
    ],
    "sm_felling_species_analysis": [
        ("species", "प्रजाति"),
        ("local_name", "स्थानीय नाम"),
        ("tree_count", "रूख गणना"),
        ("percent", "प्रतिशत"),
        ("timber_m3", "काठ घ.मी."),
        ("firewood_m3", "दाउरा घ.मी."),
        ("gross_volume_m3", "ग्रस आयतन घ.मी."),
        ("net_volume_m3", "नेट आयतन घ.मी."),
        ("avg_dbh_cm", "औसत ब्यास से.मी."),
    ],
    # --- Legend tables for chart symbolization ---
    "sm_mf_hierarchy_legend": [("symbol", "प्रतीक"), ("label", "विवरण")],
    "sm_stand_type_legend":   [("symbol", "प्रतीक"), ("label", "विवरण")],
    "sm_carbon_legend":       [("symbol", "प्रतीक"), ("label", "विवरण")],
    "sm_volume_legend":       [("symbol", "प्रतीक"), ("label", "विवरण")],
    "sm_mf_species_legend":   [("symbol", "प्रतीक"), ("label", "विवरण")],
    "sm_felling_species_legend": [("symbol", "प्रतीक"), ("label", "विवरण")],
}

def _add_activity_plan_detail_table(doc: Document, val: list):
    if not val or not isinstance(val, list):
        return
    keys = [eng_key for eng_key, _ in NP_HEADERS_ACTIVITY_PLAN if eng_key in val[0]]
    headers = [np_header for eng_key, np_header in NP_HEADERS_ACTIVITY_PLAN if eng_key in val[0]]
    widths = [NP_COL_WIDTHS_CM[i] for i, (eng_key, _) in enumerate(NP_HEADERS_ACTIVITY_PLAN) if eng_key in val[0]]
    num_cols = len(keys)
    num_rows = len(val) + 1
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    for ci, np_header in enumerate(headers):
        cell = tbl.cell(0, ci)
        cell.width = Cm(widths[ci])
        lines = np_header.split("\n")
        for li, line in enumerate(lines):
            if li == 0:
                cell.text = ""
            p = cell.add_paragraph() if li > 0 else cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "006400")
    for ri, row in enumerate(val, 1):
        for ci, key in enumerate(keys):
            cell = tbl.cell(ri, ci)
            cell.width = Cm(widths[ci])
            val_raw = row.get(key, "")
            if key == "total_budget":
                try:
                    val_raw = round(float(val_raw) / 1000)
                except (ValueError, TypeError):
                    pass
            cell.text = _fmt_value(val_raw, "ya_activity_plan_detail")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def _add_hh_records_table(doc: Document, val: list, raw_data: dict = None):
    if not val or not isinstance(val, list):
        return
    hh = (raw_data or {}).get("households", {})
    total_hh = hh.get("total_households", len(val))
    total_pop = hh.get("total_population", 0)
    forest_occ = hh.get("forest_based_occupation", 0)

    # Switch to landscape for this wide table
    landscape = doc.add_section()
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width = Cm(29.7)
    landscape.page_height = Cm(21.0)
    landscape.top_margin = Cm(1.5)
    landscape.bottom_margin = Cm(1.5)
    landscape.left_margin = Cm(1.5)
    landscape.right_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{_fix(format_devanagari(total_hh, 0))} घरपरिवारको विस्तृत विवरण")
    run.font.size = Pt(13)
    run.font.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(f"कुल जनसंख्या: {_fix(format_devanagari(total_pop, 0))}  |  वनमा आश्रित: {_fix(format_devanagari(forest_occ, 0))}")
    run2.font.size = Pt(10)
    run2.font.italic = True
    keys = [k for k, _, _, _ in NP_HEADERS_HH_RECORDS]
    headers = [h for _, h, _, _ in NP_HEADERS_HH_RECORDS]
    col_types = [t for _, _, _, t in NP_HEADERS_HH_RECORDS]
    widths = [w for _, _, w, _ in NP_HEADERS_HH_RECORDS]
    num_cols = len(keys)
    num_rows = len(val) + 1
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    for ci, np_header in enumerate(headers):
        cell = tbl.cell(0, ci)
        cell.width = Cm(widths[ci])
        lines = np_header.split("\n")
        for li, line in enumerate(lines):
            if li == 0:
                cell.text = ""
            p = cell.add_paragraph() if li > 0 else cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "006400")
        if col_types[ci] in ("number", "vertical"):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            textDirection = OxmlElement("w:textDirection")
            textDirection.set(qn("w:val"), "btLr")
            tcPr.append(textDirection)
    tbl.rows[0].height = Cm(1.4)
    tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for ri, row in enumerate(val, 1):
        for ci, key in enumerate(keys):
            cell = tbl.cell(ri, ci)
            cell.width = Cm(widths[ci])
            val_raw = row.get(key, "")
            if key == "जग्गा_एकाइ":
                val_str = "रो." if str(val_raw).strip().lower() == "ropani" else _fix(str(val_raw))
            elif key in ("घाँस_भारी", "दाउरा_भारी", "ओछ्यान_भारी"):
                val_str = _fix(format_devanagari(val_raw, 0))
            else:
                val_str = _fmt_value(val_raw, "hh_records")
            cell.text = val_str
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

    # Restore portrait for subsequent content
    portrait = doc.add_section()
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width = Cm(21.0)
    portrait.page_height = Cm(29.7)
    portrait.top_margin = Cm(2)
    portrait.bottom_margin = Cm(2)
    portrait.left_margin = Cm(2.5)
    portrait.right_margin = Cm(2.5)


# ── Field Inventory Tables ──

FI_YIELD_HEADERS = [
    ("block_name", "ब्लक", 2.5, "left"),
    ("plot_no", "प्लट नं.", 1.2, "center"),
    ("firewood", "दाउरा (के.जी.)", 2.5, "right"),
    ("grass", "घाँस (के.जी.)", 2.5, "right"),
    ("bedding", "ओछ्यान (के.जी.)", 2.5, "right"),
]

ST_TYPE_MAP = {
    "Regeneration": "पुनरुत्पादन",
    "Sapling": "लाथ्रा",
    "Pole": "पोल",
    "Tree": "रूख",
}

ST_TYPE_SORT = {"Regeneration": 0, "Sapling": 1, "Pole": 2, "Tree": 3}


def _start_two_column_section(doc: Document):
    """Insert a continuous section break within the current page, then switch to 2-column layout."""
    from docx.enum.section import WD_SECTION
    doc.add_section(WD_SECTION.CONTINUOUS)
    sect = doc.sections[-1]
    sect.top_margin = doc.sections[0].top_margin
    sect.bottom_margin = doc.sections[0].bottom_margin
    sect.left_margin = doc.sections[0].left_margin
    sect.right_margin = doc.sections[0].right_margin
    sectPr = sect._sectPr
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "480")
    cols.set(qn("w:equalWidth"), "true")
    sectPr.append(cols)


def _end_two_column_section(doc: Document):
    """Add a NEW_PAGE section break and strip any 2-column settings inherited from the prior section."""
    from docx.enum.section import WD_SECTION
    new_sect = doc.add_section(WD_SECTION.NEW_PAGE)
    new_sect.top_margin = doc.sections[0].top_margin
    new_sect.bottom_margin = doc.sections[0].bottom_margin
    new_sect.left_margin = doc.sections[0].left_margin
    new_sect.right_margin = doc.sections[0].right_margin
    # doc.add_section() COPIES all children from the old sectPr (including our 2-column w:cols).
    # Remove every w:cols element so this section defaults to single column.
    sectPr = new_sect._sectPr
    for child in list(sectPr):
        if child.tag == qn('w:cols'):
            sectPr.remove(child)


def _add_field_inventory_tables(doc: Document, calculation_id: UUID, db: Session):
    fi_calc = db.query(FieldInventoryCalculation).filter(
        FieldInventoryCalculation.calculation_id == calculation_id
    ).first()
    if not fi_calc:
        _add_no_data_para(doc)
        return

    measurements = _get_field_inventory_raw_data(db, fi_calc.id)
    if not measurements:
        _add_no_data_para(doc)
        return

    # ── Page break before this table ──
    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(_fix("वन श्रोत मापन सर्वेक्षण फारम"))
    run.font.size = Pt(11)
    run.font.bold = True

    # ── Dash-separated values in two-column monospace layout ──
    _start_two_column_section(doc)

    sep = "-"
    csv_headers = sep.join(["ब्लक", "प्लट", "प्रकार", "वैज्ञानिक नाम", "DBH", "उचाइ", "वर्ग", "गणना"])
    csv_lines = [csv_headers]

    for m in measurements:
        vals = [
            m.block_name or "",
            str(m.plot_no or ""),
            ST_TYPE_MAP.get(m.stand_type, m.stand_type or ""),
            m.scientific_name or "",
            _csv_num(m.dbh_cm),
            _csv_num(m.height_m),
            m.tree_class or "",
            str(m.count) if m.count is not None else "1",
        ]
        csv_lines.append(sep.join(vals))

    csv_text = "\n".join(csv_lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(9)
    run = p.add_run(_fix(csv_text))
    run.font.name = "Courier New"
    run.font.size = Pt(7)

    # ── Resource Yield (dash-separated, same two-column) ──
    yield_rows = _get_resource_yield(db, fi_calc.id)
    if yield_rows:
        doc.add_paragraph()
        sub_heading = doc.add_paragraph()
        sub_heading.paragraph_format.space_after = Pt(2)
        run = sub_heading.add_run(_fix("नमुना प्लटमा उपलब्ध झिजा दाउरा, सोतर तथा घाँसहरूको १०० वर्ग मि. क्षेत्रमा उपलब्धता के.जि.परिणाम"))
        run.font.size = Pt(10)
        run.font.bold = True

        y_csv = sep.join(["ब्लक", "प्लट नं.", "दाउरा", "घाँस", "सोतर"])
        y_lines = [y_csv]
        for row in yield_rows:
            vals = [
                row.get("block_name", ""),
                row.get("plot_no", ""),
                _csv_num(row.get("firewood")),
                _csv_num(row.get("grass")),
                _csv_num(row.get("bedding")),
            ]
            y_lines.append(sep.join(vals))

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(9)
        run = p.add_run(_fix("\n".join(y_lines)))
        run.font.name = "Courier New"
        run.font.size = Pt(7)

    _end_two_column_section(doc)
    doc.add_paragraph()


def _add_no_data_para(doc: Document):
    no_data = doc.add_paragraph()
    run = no_data.add_run("[Table data not available: fieldinventory]")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(200, 50, 50)


def _csv_num(val) -> str:
    """Format a numeric value for CSV — empty string if None, Devanagari digits."""
    if val is None:
        return ""
    return format_devanagari(val, 1) if isinstance(val, (int, float)) else str(val)


def _get_field_inventory_raw_data(db: Session, fi_calc_id: int) -> list:
    """Fetch raw measurement rows — only input columns, no calculated fields.
    Uses raw SQL with namedtuple-like rows (no ORM overhead)."""
    from sqlalchemy import text as sa_text
    rows = db.execute(sa_text("""
        SELECT
            sp.block_name,
            sp.sample_plot_number AS plot_no,
            m.stand_type,
            m.species_local AS local_name,
            m.species_scientific AS scientific_name,
            m.dbh_cm,
            m.height_m,
            m.tree_class,
            m.count,
            m.sn
        FROM field_inventory_measurements m
        JOIN field_inventory_sample_plots sp ON sp.id = m.sample_plot_id
        WHERE sp.field_inventory_calculation_id = :fi_calc_id
        ORDER BY sp.block_name, sp.sample_plot_number, m.stand_type, m.sn
    """), {"fi_calc_id": fi_calc_id}).fetchall()
    return rows


def _get_resource_yield(db: Session, fi_calc_id: int) -> list:
    """Return resource yield per plot — queried directly from sample plots."""
    from sqlalchemy import text as sa_text
    rows = db.execute(sa_text("""
        SELECT
            block_name,
            sample_plot_number AS plot_no,
            firewood_kg_per_100sqm_per_year AS firewood,
            grass_kg_per_100sqm_per_year AS grass,
            bedding_material_kg_per_100sqm_per_year AS bedding
        FROM field_inventory_sample_plots
        WHERE field_inventory_calculation_id = :fi_calc_id
        ORDER BY block_name, sample_plot_number
    """), {"fi_calc_id": fi_calc_id}).fetchall()
    return [
        {
            "block_name": r.block_name,
            "plot_no": str(r.plot_no),
            "firewood": r.firewood,
            "grass": r.grass,
            "bedding": r.bedding,
        }
        for r in rows
    ]


def _add_text_content(doc: Document, text: str, calculation_id: UUID = None, db: Session = None, raw_data: dict = None, table_cache: dict = None):
    text = _fix( text)
    parts = re.split(r"(\{\{chart:\w+\}\}|\{\{map:\w+\}\}|\{\{table:\w+\}\})", text)
    for part in parts:
        part = part.strip()
        if not part:
            continue

        chart_match = re.match(r"\{\{chart:(\w+)\}\}", part)
        map_match = re.match(r"\{\{map:(\w+)\}\}", part)
        table_match = re.match(r"\{\{table:(\w+)\}\}", part)

        if chart_match and raw_data:
            _add_chart_from_type(doc, chart_match.group(1), raw_data, calculation_id)
        elif map_match and calculation_id and db:
            _add_map_standard(doc, map_match.group(1), calculation_id, db)
        elif table_match and calculation_id:
            table_id = table_match.group(1)
            if table_id == "fieldinventory":
                _add_field_inventory_tables(doc, calculation_id, db)
            else:
                _add_table_inline(doc, table_id, table_cache)
        else:
            var_match = re.match(r"^\{\{(\w+(?::\w+)*)\}\}$", part)
            if var_match and raw_data:
                var_name = var_match.group(1)
                if var_name.endswith(":full"):
                    _add_section_full_docx(doc, var_name, raw_data, calculation_id)
                    continue
                if not var_name.startswith("chart:") and not var_name.startswith("map:") and not var_name.startswith("table:"):
                    var_val = _resolve_var_from_raw(var_name, raw_data)
                    if isinstance(var_val, list):
                        if var_name == "uc_members":
                            _add_uc_members_table(doc, var_val)
                            continue
                        if var_name == "ya_activity_plan_detail":
                            _add_activity_plan_detail_table(doc, var_val)
                            continue
                        if var_name == "hh_records":
                            _add_hh_records_table(doc, var_val, raw_data)
                            continue
                        _add_list_table(doc, var_val, var_name)
                        continue
                    if var_val is not None and not isinstance(var_val, (dict, list)):
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run(_fmt_value(var_val, var_name))
                        run.font.size = Pt(11)
                        continue

            for para_text in part.split("\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                line_var_match = re.match(r"^\{\{(\w+(?::\w+)*)\}\}$", para_text)
                if line_var_match and raw_data:
                    var_name = line_var_match.group(1)
                    if var_name.endswith(":full"):
                        _add_section_full_docx(doc, var_name, raw_data, calculation_id)
                        continue
                    if not var_name.startswith("chart:") and not var_name.startswith("map:") and not var_name.startswith("table:"):
                        var_val = _resolve_var_from_raw(var_name, raw_data)
                        if isinstance(var_val, list):
                            if var_name == "uc_members":
                                _add_uc_members_table(doc, var_val)
                                continue
                            if var_name == "hh_records":
                                _add_hh_records_table(doc, var_val, raw_data)
                                continue
                            _add_list_table(doc, var_val, var_name)
                            continue
                        if var_val is not None and not isinstance(var_val, (dict, list)):
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(6)
                            run = p.add_run(_fmt_value(var_val, var_name))
                            run.font.size = Pt(11)
                            continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                if raw_data:
                    para_text = re.sub(r'\{\{(\w+(?::\w+)*)\}\}', lambda m: _resolve_var_text(m.group(0), raw_data) if not m.group(1).startswith(('chart:', 'map:', 'table:')) else m.group(0), para_text)
                run = p.add_run(_fix( para_text))
                run.font.size = Pt(11)
                run.font.name = "Nirmala UI"


def _add_section_full_docx(doc: Document, var_name: str, raw_data: dict, calculation_id: UUID = None):
    section_name = var_name.replace("section:", "").replace(":full", "")
    key = f"section:{section_name}"
    sections = raw_data.get("section_generators", {})
    narrative = sections.get(key, "")
    if not narrative:
        return
    title_np, title_en = _SECTION_TITLES.get(section_name, (section_name, section_name))
    heading = doc.add_heading(_fix(title_np), level=2)
    for r in heading.runs:
        r.font.color.rgb = _COVER_GREEN
    if title_en != title_np:
        sub = doc.add_paragraph()
        run = sub.add_run(_fix(title_en))
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(120, 120, 120)
        sub.paragraph_format.space_after = Pt(6)
    for para_text in narrative.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(_fix(para_text))
        run.font.size = Pt(11)
        run.font.name = "Nirmala UI"
    chart_type = _SECTION_CHARTS.get(section_name)
    if chart_type:
        _add_chart_from_type(doc, chart_type, raw_data, calculation_id)


def _add_table_inline(doc: Document, table_id: str, table_cache: dict = None):
    from app.services.operational_plan.variable_registry import TABLE_ID_ALIAS
    table_id = TABLE_ID_ALIAS.get(table_id, table_id)
    table_data = (table_cache or {}).get(table_id)

    if not table_data or not table_data.rows:
        p = doc.add_paragraph()
        run = p.add_run(f"[Table data not available: {table_id}]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(200, 0, 0)
        return

    vdef = get_variable(f"table:{table_id}")
    title_ne = vdef.label_ne if vdef else table_id.replace("_", " ").title()
    p_heading = doc.add_paragraph()
    p_heading.paragraph_format.space_after = Pt(2)
    run_heading = p_heading.add_run(title_ne)
    run_heading.font.size = Pt(13)
    run_heading.font.bold = True

    rows = table_data.rows
    headers = list(rows[0].keys()) if rows and isinstance(rows[0], dict) else []
    data_rows = [[_fix(format_devanagari(r.get(h, ""))) for h in headers] for r in rows] if headers else rows
    num_cols = len(headers) if headers else len(rows[0]) if rows else 1
    num_rows = len(rows) + 1

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    np_headers = NP_HEADERS_BIODIVERSITY.get(table_id, [])
    np_map = {eng: np for eng, np in np_headers}
    for ci, header in enumerate(headers):
        cell = tbl.cell(0, ci)
        cell.text = np_map.get(header, header.replace("_", " ").title())
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "006400")

    for ri, row in enumerate(data_rows, 1):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()


def _short_hierarchy(r: dict) -> str:
    sub = (r.get("sub_compartment") or "-").strip()
    comp = (r.get("compartment") or "-").strip()
    label = sub if sub and sub != "-" else comp
    return label if label and label != "-" else "-"


def _render_sm_chart_internal(chart_type: str, raw_data: dict, forest_name: str = "") -> Optional[str]:
    """Render a tree-mapping sm_* chart. Returns SVG data URI or None."""
    sm = raw_data.get("tree_mapping_analysis", {})
    if chart_type == "sm_felling_dbh_pie":
        fda = sm.get("sm_felling_dbh_analysis", [])
        if fda:
            labels = [r.get("dbh_class", str(i)) for i, r in enumerate(fda)]
            values = [r.get("tree_count", 0) for r in fda]
            pcts = [r.get("percent", 0) for r in fda]
            return _chart_from_data(labels, values, forest_name, "DBH वर्ग अनुसार कटानी रूख", percentages=pcts)
    elif chart_type == "sm_felling_species_bar":
        fsa = sm.get("sm_felling_species_analysis", [])[:10]
        if fsa:
            labels = [format_devanagari(i, 0) for i in range(1, len(fsa) + 1)]
            values = [r.get("tree_count", 0) for r in fsa]
            pcts = [r.get("percent", 0) for r in fsa]
            return _chart_from_data(labels, values, forest_name, "प्रजाति अनुसार कटानी रूख", is_pie=False, percentages=pcts)
    elif chart_type == "sm_mother_felling_pie":
        mfs = sm.get("sm_mother_felling_summary", {})
        if mfs:
            labels = ["माँउ रूख", "कटानी रूख"]
            values = [mfs.get("total_mother_trees", 0), mfs.get("total_felling_trees", 0)]
            return _chart_from_data(labels, values, forest_name, "माँउ रूख बनाम कटानी रूख", colors=["#22c55e", "#ef4444"])
    elif chart_type == "sm_mother_felling_species_bar":
        mbs = sm.get("sm_mother_tree_by_species", [])[:10]
        fbs = sm.get("sm_felling_tree_by_species", [])[:10]
        if mbs or fbs:
            all_species = list(dict.fromkeys([r.get("species", "") for r in mbs] + [r.get("species", "") for r in fbs]))[:10]
            labels = [format_devanagari(i, 0) for i in range(1, len(all_species) + 1)]
            mother_map = {r.get("species", ""): r.get("tree_count", 0) for r in mbs}
            felling_map = {r.get("species", ""): r.get("tree_count", 0) for r in fbs}
            mother_vals = [mother_map.get(sp, 0) for sp in all_species]
            felling_vals = [felling_map.get(sp, 0) for sp in all_species]
            return _chart_from_data_grouped(
                labels,
                {"माँउ": mother_vals, "कटानी": felling_vals},
                forest_name, "प्रजाति अनुसार माँउ बनाम कटानी",
                colors=["#22c55e", "#ef4444"],
            )
    elif chart_type == "sm_stand_type_bar":
        st = sm.get("sm_stand_type_by_hierarchy", [])
        if st:
            labels_short = [format_devanagari(i, 0) for i in range(1, len(st) + 1)]
            regen = [r.get("regeneration", 0) for r in st]
            sapling = [r.get("sapling", 0) for r in st]
            pole = [r.get("pole", 0) for r in st]
            tree = [r.get("tree", 0) for r in st]
            return _chart_from_data_grouped(
                labels_short,
                {"पुनरुत्पादन": regen, "लाथ्रा": sapling, "पोल": pole, "रूख": tree},
                forest_name, "स्तर अनुसार स्ट्यान्ड प्रकार",
                colors=["#22c55e", "#3b82f6", "#f59e0b", "#ef4444"],
            )
    elif chart_type == "sm_carbon_bar":
        cb = sm.get("sm_carbon_by_hierarchy", [])
        if cb:
            labels_short = [format_devanagari(i, 0) for i in range(1, len(cb) + 1)]
            carbon_vals = [r.get("carbon_tc", 0) for r in cb]
            co2_vals = [r.get("co2_tco2", 0) for r in cb]
            return _chart_from_data_grouped(
                labels_short,
                {"कार्बन (tC)": carbon_vals, "CO₂e (tCO₂)": co2_vals},
                forest_name, "स्तर अनुसार कार्बन मौज्दात",
                colors=["#22c55e", "#3b82f6"],
            )
    elif chart_type == "sm_volume_bar":
        vb = sm.get("sm_volume_by_hierarchy", [])
        if vb:
            labels_short = [format_devanagari(i, 0) for i in range(1, len(vb) + 1)]
            stem = [r.get("stem_volume_m3", 0) for r in vb]
            branch = [r.get("branch_volume_m3", 0) for r in vb]
            return _chart_from_data_grouped(
                labels_short,
                {"काण्ड": stem, "हाँगा": branch},
                forest_name, "स्तर अनुसार आयतन संरचना",
                colors=["#22c55e", "#3b82f6"],
            )
    elif chart_type == "sm_mother_felling_hierarchy_bar":
        hs = sm.get("sm_hierarchy_summary", [])
        rbd = sm.get("sm_hierarchy_remark_breakdown", {})
        if hs:
            labels_short = [format_devanagari(i, 0) for i in range(1, len(hs) + 1)]
            mother_vals = []
            felling_vals = []
            for r in hs:
                key = f"{r.get('compartment','-')}|{r.get('sub_compartment','-')}"
                breakdown = rbd.get(key, {})
                mother_vals.append(breakdown.get("mother_trees", 0))
                felling_vals.append(breakdown.get("felling_trees", 0))
            return _chart_from_data_grouped(
                labels_short,
                {"माँउ": mother_vals, "कटानी": felling_vals},
                forest_name, "माँउ बनाम कटानी",
                colors=["#22c55e", "#ef4444"],
            )
    return None


def _add_chart_from_type(doc: Document, chart_type: str, raw_data: dict, calculation_id: UUID = None):
    # Auto-append legend for sm_* hierarchy charts
    _SM_LEGEND_MAP = {
        "sm_mother_felling_hierarchy_bar": "sm_mf_hierarchy_legend",
        "sm_stand_type_bar": "sm_stand_type_legend",
        "sm_carbon_bar": "sm_carbon_legend",
        "sm_volume_bar": "sm_volume_legend",
        "sm_mother_felling_species_bar": "sm_mf_species_legend",
        "sm_felling_species_bar": "sm_felling_species_legend",
    }
    def _append_sm_legend(doc, chart_type, raw_data):
        legend_var = _SM_LEGEND_MAP.get(chart_type)
        if legend_var and raw_data:
            legend_data = _resolve_var_from_raw(legend_var, raw_data)
            if isinstance(legend_data, list) and legend_data:
                _add_list_table(doc, legend_data, legend_var)

    # Check cache first
    if calculation_id:
        cached = _chart_cache_get(calculation_id, chart_type)
        if cached:
            from app.utils.svg_to_png import add_svg_picture
            add_svg_picture(doc, cached.getvalue(), width_inches=5.0)
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(chart_type.replace("_", " ").title())
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()
            _append_sm_legend(doc, chart_type, raw_data)
            return

    forest_name = raw_data.get("basic_info", {}).get("forest_name", "")
    img_data = _generate_chart(chart_type, raw_data, forest_name)

    if img_data:
        try:
            if img_data.startswith("data:"):
                from app.utils.svg_to_png import add_svg_picture, _decode_data_uri
                add_svg_picture(doc, img_data, width_inches=5.0)
                if calculation_id:
                    raw_bytes, _ = _decode_data_uri(img_data)
                    _chart_cache_set(calculation_id, chart_type, BytesIO(raw_bytes))
            else:
                with open(img_data, "rb") as _f:
                    img_bytes = _f.read()
                from app.utils.svg_to_png import add_svg_picture
                add_svg_picture(doc, img_bytes, width_inches=5.0)
                if calculation_id:
                    _chart_cache_set(calculation_id, chart_type, BytesIO(img_bytes))
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(chart_type.replace("_", " ").title())
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()
            _append_sm_legend(doc, chart_type, raw_data)
            return
        except Exception:
            pass

    p = doc.add_paragraph()
    run = p.add_run(f"[Chart data not available: {chart_type}]")
    run.font.italic = True
    run.font.color.rgb = RGBColor(200, 0, 0)


def _dev_val(v):
    """Format a value with Devanagari digits + %."""
    return format_devanagari(v, 1) + "%"

def _chart_from_data(labels, values, forest_name, title, is_pie=True, colors=None, legend_cols=3, percentages=None):
    _ensure_dev_font()
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    fig, ax = plt.subplots(figsize=(6, 4), dpi=150)
    if is_pie:
        if colors is None:
            colors = plt.cm.Set3(np.linspace(0, 1, len(labels)))
        wedges, texts = ax.pie(
            values, labels=None, startangle=90,
            colors=colors,
        )
        legend_labels = [f"{l} ({_dev_val(v)})" for l, v in zip(labels, values)]
        ax.legend(
            wedges, legend_labels, loc='lower center',
            bbox_to_anchor=(0.5, -0.18), ncol=min(legend_cols, len(labels)),
            fontsize=7, frameon=False, prop=_dev_fontprop(7),
        )
        ax.set_title(f'{forest_name} - {title}', fontsize=11, fontweight='bold', pad=15)
    else:
        if colors is None:
            colors = ['#2ecc71', '#3498db', '#e67e22', '#e74c3c', '#9b59b6', '#f1c40f']
        bars = ax.bar(range(len(labels)), values, color=colors[:len(labels)], edgecolor='white')
        for i, (bar, val) in enumerate(zip(bars, values)):
            if percentages and i < len(percentages):
                label_text = f"{format_devanagari(val, 1)} ({format_devanagari(percentages[i], 1)}%)"
            else:
                label_text = format_devanagari(val, 1)
            ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + max(values)*0.01,
                    label_text, ha='center', va='bottom', fontsize=7,
                    fontproperties=_dev_fontprop(7))
        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=8)
        ax.set_ylabel('Value', fontsize=10)
        ax.set_title(f'{forest_name} - {title}', fontsize=11, fontweight='bold')
        ax.yaxis.grid(True, alpha=0.3)

    if _DEV_FONT_PATH:
        ax.title.set_fontproperties(_dev_fontprop(11))
        ax.xaxis.label.set_fontproperties(_dev_fontprop(10))
        ax.yaxis.label.set_fontproperties(_dev_fontprop(10))
        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(_dev_fontprop(8))

    fig.tight_layout()
    if is_pie:
        fig.subplots_adjust(bottom=0.28)
    buf = BytesIO()
    fig.savefig(buf, format='svg', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    from base64 import b64encode
    return f"data:image/svg+xml;base64,{b64encode(buf.read()).decode()}"


def _chart_from_data_grouped(labels, series_dict, forest_name, title, colors=None):
    """Create a grouped bar chart with multiple series (e.g. demand vs supply)."""
    _ensure_dev_font()
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np

    series_names = list(series_dict.keys())
    series_values = list(series_dict.values())
    n_series = len(series_names)

    if colors is None:
        colors = ['#dc2626', '#059669', '#3498db', '#e67e22', '#9b59b6']
    colors = colors[:n_series]

    fig, ax = plt.subplots(figsize=(8, 4.5), dpi=150)

    x = np.arange(len(labels))
    width = 0.8 / max(n_series, 1)
    max_val = max(max(v) for v in series_values) if series_values else 0

    for i, (sname, svals) in enumerate(series_dict.items()):
        offset = (i - (n_series - 1) / 2) * width
        bars = ax.bar(x + offset, svals, width, label=sname, color=colors[i], edgecolor='white')
        for bar, val in zip(bars, svals):
            if val:
                ax.text(bar.get_x() + bar.get_width() / 2.,
                        bar.get_height() + max_val * 0.01,
                        format_devanagari(val, 1),
                        ha='center', va='bottom', fontsize=6,
                        fontproperties=_dev_fontprop(6))

    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=8)
    ax.set_ylabel('Value', fontsize=10)
    ax.set_title(f'{forest_name} - {title}', fontsize=11, fontweight='bold')
    ax.legend(fontsize=8, prop=_dev_fontprop(8))
    ax.yaxis.grid(True, alpha=0.3)

    if _DEV_FONT_PATH:
        ax.title.set_fontproperties(_dev_fontprop(11))
        ax.xaxis.label.set_fontproperties(_dev_fontprop(10))
        ax.yaxis.label.set_fontproperties(_dev_fontprop(10))
        for label in ax.get_xticklabels() + ax.get_yticklabels():
            label.set_fontproperties(_dev_fontprop(8))

    fig.tight_layout()
    buf = BytesIO()
    fig.savefig(buf, format='svg', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    from base64 import b64encode
    return f"data:image/svg+xml;base64,{b64encode(buf.read()).decode()}"


def _add_map_standard(doc: Document, map_type: str, calculation_id: UUID, db: Session, forest_name: str = "CF"):
    from app.services.management_plan_docx.plan_map_service import generate_standard_map, LAYER_LABELS

    alias_map = {"boundary_map": "boundary"}
    layer_name = alias_map.get(map_type, map_type)
    known_layers = {"boundary","forest_type","forest_health","slope","biomass","landcover","soil_texture","dem","aspect","canopy","sampling_plot","sampling_plot_topo","sampling_plot_satellite","fieldbook","usergroup","subarea","compartment","sub_compartment"}

    if layer_name not in known_layers:
        p = doc.add_paragraph()
        run = p.add_run(f"[Unknown map type: {map_type}]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(200, 0, 0)
        return

    if layer_name == "subarea":
        from app.models.calculation import Calculation
        calc = db.query(Calculation).filter(Calculation.id == calculation_id).first()
        has_sub_areas = bool(
            calc and calc.result_data and calc.result_data.get("sub_areas")
        )
        if not has_sub_areas:
            p = doc.add_paragraph()
            run = p.add_run(f"[Map: {map_type}] - data not available (no sub-areas drawn)")
            run.font.italic = True
            run.font.color.rgb = RGBColor(200, 0, 0)
            return

    if layer_name in ("compartment", "sub_compartment"):
        from app.models.forest_block import ForestBlock
        has_compartments = db.query(ForestBlock.id).filter(
            ForestBlock.calculation_id == str(calculation_id),
            ForestBlock.is_compartment == True,
        ).first() is not None
        if not has_compartments:
            p = doc.add_paragraph()
            run = p.add_run(f"[Map: {map_type}] - data not available (no compartments drawn)")
            run.font.italic = True
            run.font.color.rgb = RGBColor(200, 0, 0)
            return

    img_size = 0
    try:
        use_cache = True
        buf = generate_standard_map(db, calculation_id, layer_name, forest_name=forest_name, use_cache=use_cache)
        buf.seek(0, 2)
        img_size = buf.tell()
        buf.seek(0)
        if img_size < 5000:
            logger.warning(f"Map too small ({img_size}B) for {layer_name} — generation likely failed")
            buf = None
        else:
            logger.info(f"Map OK ({img_size}B) — adding {layer_name} to DOCX")
    except Exception as e:
        logger.error(f"Map generation exception for {layer_name}: {e}")
        buf = None

    if buf:
        doc.add_picture(buf, width=Inches(5.83 if layer_name == "subarea" else 5.5))
        labels = LAYER_LABELS.get(layer_name, {"ne": layer_name, "en": layer_name})
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_p.add_run(labels.get("en", layer_name))
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[Map: {map_type}] - map unavailable (PNG: {img_size}B)")
        run.font.italic = True
        run.font.color.rgb = RGBColor(200, 0, 0)


def _generate_chart(chart_type: str, raw_data: dict, forest_name: str = "") -> str | None:
    """Shared chart generation: returns img_data (data-URI string or file path) or None.

    This function is the single source of truth for which chart_type maps to
    which generator call.  Both ``_add_chart`` (DOCX) and ``_render_chart_html``
    delegate the dispatch here and only handle the final rendering step.
    """
    from app.services.report.chart_generator import (
        generate_species_pie, generate_forest_type_pie, generate_block_area_bar,
        generate_dbh_histogram, generate_biomass_bar, generate_slope_pie,
        generate_canopy_pie, generate_landcover_pie, generate_ug_landcover_pie,
    )

    language = raw_data.get("basic_info", {}).get("language", "NP")

    # ---- helper lambdas (only for truly trivial extractions) ----
    def _ra(section: str) -> dict:
        return raw_data.get("raster_analysis", {}).get(section, {})

    img_data = None

    # ---------- species / forest ----------
    if chart_type.startswith("sm_"):
        img_data = _render_sm_chart_internal(chart_type, raw_data, forest_name)
    elif chart_type in ("species_pie", "species_composition_pie"):
        species = raw_data.get("species", {})
        if isinstance(species, dict):
            species = species.get("species_list", [])
        if isinstance(species, dict):
            species = species.get("species_list", [])
        if species:
            img_data = generate_species_pie(species, forest_name, top_n=8)
    elif chart_type == "species_composition_pie_fi":
        fi = raw_data.get("field_inventory", {})
        comp = fi.get("fi_species_composition", {})
        if comp and isinstance(comp, dict):
            species_list = [
                {"scientific_name": k, "local_name": "", "availability_rank": i}
                for i, (k, _) in enumerate(
                    sorted(comp.items(), key=lambda x: x[1], reverse=True)
                )
            ]
            img_data = generate_species_pie(species_list, forest_name, top_n=8)
    elif chart_type in ("forest_type_pie", "forest_type"):
        ft = _ra("forest_type").get("percentages", {})
        img_data = generate_forest_type_pie(ft, forest_name, language=language)
    elif chart_type == "block_area_bar":
        blocks = raw_data.get("blocks", {}).get("blocks", [])
        img_data = generate_block_area_bar(blocks, forest_name, language=language)

    # ---------- inventory ----------
    elif chart_type == "dbh_histogram":
        inv = raw_data.get("inventory", {})
        dbh = inv.get("dbh_summary", {}) or inv.get("dbh_distribution", {})
        img_data = generate_dbh_histogram(dbh, forest_name)
    elif chart_type == "biomass_bar":
        bi = raw_data.get("basic_info", {})
        agb = bi.get("above_ground_biomass_tons", 0) or bi.get("agb_total", 0)
        carbon = bi.get("carbon_stock_tc", 0) or bi.get("carbon_stock", 0)
        img_data = generate_biomass_bar(agb, carbon, forest_name, language=language)

    # ---------- raster analysis pies ----------
    elif chart_type in ("slope_pie", "slope_bar"):
        sp = _ra("slope").get("percentages", {})
        dom = _ra("slope").get("dominant_class", "")
        img_data = generate_slope_pie(sp, dom, forest_name, language=language)
    elif chart_type in ("canopy_pie", "canopy_bar"):
        cp = _ra("canopy").get("percentages", {})
        dom = _ra("canopy").get("dominant_class", "")
        img_data = generate_canopy_pie(cp, dom, forest_name, language=language)
    elif chart_type == "landcover_pie":
        lc = _ra("landcover").get("percentages", {})
        dom = _ra("landcover").get("dominant_class", "")
        img_data = generate_landcover_pie(lc, dom, forest_name, language=language)
    elif chart_type == "forest_health_pie":
        fh = _ra("forest_health").get("percentages", {})
        if fh:
            fh_colors = [_HEALTH_COLORS_MAP.get(k, "#95a5a6") for k in fh.keys()]
            img_data = _chart_from_data(list(fh.keys()), list(fh.values()), forest_name, "Forest Health", colors=fh_colors, legend_cols=2)
    elif chart_type == "aspect_rose":
        ap = _ra("aspect").get("percentages", {})
        if ap:
            asp_colors = [_ASPECT_COLORS_MAP.get(k, "#95a5a6") for k in ap.keys()]
            img_data = _chart_from_data(list(ap.keys()), list(ap.values()), forest_name, "Aspect Distribution", colors=asp_colors, legend_cols=4)
    elif chart_type == "nasa_forest_2020_pie":
        rd = raw_data.get("result_data", {})
        pct = rd.get("whole_nasa_forest_2020_percentages", {})
        if pct and isinstance(pct, dict):
            items = {k: v for k, v in pct.items() if v > 0}
            if items:
                nasa_colors = [_NASA_FOREST_QUALITY_COLORS.get(k, "#95a5a6") for k in items.keys()]
                img_data = _chart_from_data(list(items.keys()), list(items.values()), forest_name, "Forest Quality", colors=nasa_colors)
    elif chart_type == "soil_bar":
        sp = _ra("soil").get("percentages", {})
        if not sp:
            rd = raw_data.get("result_data", {})
            props = rd.get("soil_properties", {})
            if props and props.get("clay_pct") is not None:
                sp = {"Clay": props["clay_pct"], "Sand": props["sand_pct"], "Silt": props["silt_pct"]}
        if sp:
            img_data = _chart_from_data(list(sp.keys()), list(sp.values()), forest_name, "Soil Distribution", is_pie=False, colors=_SOIL_BAR_COLORS[:len(sp)])

    # ---------- user group / households ----------
    elif chart_type == "ug_land_cover_classes_chart":
        ug = raw_data.get("user_group", {})
        lc_classes = ug.get("land_cover_classes", [])
        if lc_classes:
            img_data = generate_ug_landcover_pie(lc_classes, forest_name, language=language)
    elif chart_type == "hh_prosperity_pie":
        hh = raw_data.get("households", {}).get("prosperity_distribution", {})
        if hh and isinstance(hh, dict):
            img_data = _chart_from_data(list(hh.keys()), list(hh.values()), forest_name, "Prosperity Distribution")
    elif chart_type == "hh_caste_bar":
        hh = raw_data.get("households", {}).get("caste_distribution", {})
        if hh and isinstance(hh, dict):
            img_data = _chart_from_data(list(hh.keys()), list(hh.values()), forest_name, "Caste Distribution", is_pie=False)
    elif chart_type == "hh_caste_pie":
        hh = raw_data.get("households", {}).get("caste_distribution", {})
        if hh and isinstance(hh, dict):
            img_data = _chart_from_data(list(hh.keys()), list(hh.values()), forest_name, "जाति वितरण (Caste Distribution)")
    elif chart_type == "hh_prosperity_bar":
        hh = raw_data.get("households", {}).get("prosperity_distribution", {})
        if hh and isinstance(hh, dict):
            img_data = _chart_from_data(list(hh.keys()), list(hh.values()), forest_name, "समृद्धि वितरण (Prosperity Distribution)", is_pie=False)

    # ---------- demand & supply ----------
    elif chart_type in ("hh_demand_supply_bar", "demand_supply_bar", "demand_supply_deficit_bar"):
        ds = raw_data.get("demand_supply", {})
        if ds and ds.get("demand"):
            products = ["firewood_bhari", "grass_bhari", "bedding_bhari", "timber_cft", "poles_count"]
            labels = [_DS_PRODUCT_LABELS.get(k, k) for k in products]
            if chart_type == "hh_demand_supply_bar":
                demand_vals = [ds.get("demand", {}).get(k, 0) or 0 for k in products]
                supply_vals = [ds.get("total_supply", {}).get(k, 0) or 0 for k in products]
                img_data = _chart_from_data_grouped(
                    labels, {"माग": demand_vals, "आपूर्ति": supply_vals},
                    forest_name, "माग र आपूर्ति तुलना (Demand & Supply Comparison)",
                    colors=["#dc2626", "#059669"],
                )
            elif chart_type == "demand_supply_bar":
                demand_vals = [ds.get("demand", {}).get(k, 0) or 0 for k in products]
                cf_reg = [ds.get("supply_cf_regular", {}).get(k, 0) or 0 for k in products]
                cf_aah = [ds.get("supply_cf_aah", {}).get(k, 0) or 0 for k in products]
                private = [ds.get("supply_private", {}).get(k, 0) or 0 for k in products]
                total = [ds.get("total_supply", {}).get(k, 0) or 0 for k in products]
                img_data = _chart_from_data_grouped(
                    labels, {"माग": demand_vals, "सा.वन नियमित": cf_reg, "सा.वन AAH": cf_aah, "निजि क्षेत्र": private, "जम्मा आपूर्ति": total},
                    forest_name, "माग र आपूर्ति ब्रेकडाउन (Demand & Supply Breakdown)",
                    colors=["#dc2626", "#059669", "#3498db", "#e67e22", "#9b59b6"],
                )
            elif chart_type == "demand_supply_deficit_bar":
                deficit = [(ds.get("total_supply", {}).get(k, 0) or 0) - (ds.get("demand", {}).get(k, 0) or 0) for k in products]
                d_colors = ["#059669" if v >= 0 else "#dc2626" for v in deficit]
                img_data = _chart_from_data(labels, deficit, forest_name, "बचत/कमी तुलना (Surplus/Deficit Comparison)", is_pie=False, colors=d_colors)

    # ---------- activities / budget ----------
    elif chart_type == "budget_bar":
        acts = raw_data.get("activities", {})
        activities = acts.get("activities", [])
        if activities:
            labels = [f"Activity {a.get('activity_id', i + 1)}" for i, a in enumerate(activities)]
            values = [a.get("default_quantity", 0) or sum(yd.get("budget", 0) for yd in a.get("yearly_details", [])) for a in activities]
            img_data = _chart_from_data(labels, values, forest_name, "Budget", is_pie=False)
    elif chart_type == "ya_budget_year_bar":
        yp = raw_data.get("yearly_plan", {})
        trend = yp.get("budget_year_trend", {})
        if trend and isinstance(trend, dict):
            labels = [f"Year {k}" for k in sorted(trend.keys(), key=int)]
            values = [trend[k] for k in sorted(trend.keys(), key=int)]
            year_colors = ["#2ecc71", "#27ae60", "#1abc9c", "#16a085", "#3498db", "#2980b9", "#9b59b6", "#8e44ad", "#e67e22", "#d35400"]
            img_data = _chart_from_data(labels, values, forest_name, "वार्षिक बजेट वितरण (Year-wise Budget)", is_pie=False, colors=year_colors[:len(labels)])
    elif chart_type == "ya_program_pie":
        yp = raw_data.get("yearly_plan", {})
        pie_data = yp.get("program_pie_data", {})
        if pie_data and isinstance(pie_data, dict):
            prog_items = {k: v for k, v in pie_data.items() if v > 0}
            if prog_items:
                prog_colors = ["#27ae60", "#2980b9", "#e67e22", "#e74c3c", "#9b59b6", "#f1c40f", "#1abc9c", "#2c3e50"]
                img_data = _chart_from_data(list(prog_items.keys()), list(prog_items.values()), forest_name, "कार्यक्रम अनुसार बजेट (Program Budget)", colors=prog_colors[:len(prog_items)])

    # ---------- field inventory DBH bars ----------
    elif chart_type in ("dbh_class_bar", "dbh_class_count_bar"):
        cd = raw_data.get("field_inventory", {}).get("fi_dbh_class_chart_data", [])
        if cd:
            labels = [d["label"] for d in cd]
            values = [d["count_per_ha"] for d in cd]
            total = sum(values)
            pcts = [v / total * 100 if total > 0 else 0 for v in values]
            dbh_colors = ["#1a6e34", "#2d8f4e", "#45b068", "#6fc48a", "#99d8ae", "#c2ebd0"]
            img_data = _chart_from_data(labels, values, forest_name, "ब्यास क्लास अनुसार रूख संख्या (संख्या/हे.)", is_pie=False, colors=dbh_colors[:len(labels)], percentages=pcts)

    return img_data


def _add_chart(doc: Document, node: TreeNode, raw_data: Dict[str, Any], calculation_id: UUID = None):
    chart_type = node.chart_type
    if not chart_type:
        return

    forest_name = raw_data.get("basic_info", {}).get("forest_name", "")

    # Check cache first
    if calculation_id:
        cached = _chart_cache_get(calculation_id, chart_type)
        if cached:
            from app.utils.svg_to_png import add_svg_picture
            add_svg_picture(doc, cached.getvalue(), width_inches=5.5)
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(node.title_ne or "Chart")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()
            return

    img_data = _generate_chart(chart_type, raw_data, forest_name)

    if img_data:
        try:
            if img_data.startswith("data:"):
                from app.utils.svg_to_png import add_svg_picture, _decode_data_uri
                add_svg_picture(doc, img_data, width_inches=5.5)
                if calculation_id:
                    raw_bytes, _ = _decode_data_uri(img_data)
                    _chart_cache_set(calculation_id, chart_type, BytesIO(raw_bytes))
            else:
                with open(img_data, "rb") as _f:
                    img_bytes = _f.read()
                from app.utils.svg_to_png import add_svg_picture
                add_svg_picture(doc, img_bytes, width_inches=5.5)
                if calculation_id:
                    _chart_cache_set(calculation_id, chart_type, BytesIO(img_bytes))
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(node.title_ne or "Chart")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()
            return
        except Exception:
            pass

    p = doc.add_paragraph()
    run = p.add_run(f"[Chart: {node.title_ne}] - no data available")
    run.font.italic = True
    run.font.color.rgb = RGBColor(200, 0, 0)


def _add_table(doc: Document, node: TreeNode, table_cache: dict = None):
    table_id = node.table_id
    if not table_id:
        return

    table_data = (table_cache or {}).get(table_id)

    if not table_data or not table_data.rows:
        p = doc.add_paragraph()
        run = p.add_run(f"[Table: {table_id} — no data]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(150, 150, 150)
        return

    rows = table_data.rows
    if not rows:
        return

    headers = list(rows[0].keys()) if isinstance(rows[0], dict) else []
    data_rows = [[format_devanagari(r.get(h, "")) for h in headers] for r in rows] if headers else rows

    tbl = doc.add_table(rows=1 + len(data_rows), cols=len(headers) if headers else len(data_rows[0]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "006400")

    for ri, row in enumerate(data_rows, 1):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(node.title_ne or table_id)
    run.font.size = Pt(9)
    run.font.italic = True


def _build_uc_members_data(val: list) -> tuple:
    """Build uc_members table data with Nepali headers, chairperson-first layout."""
    if not val or not isinstance(val, list):
        return [], []
    headers = ["पदाधिकारीको नाम", "पद", "ठेगाना", "दाँया", "बाँया", "हस्ताक्षर"]
    chairperson = None
    others = []
    for m in val:
        pos = (m.get("position") or "").strip()
        if pos == "अध्यक्ष":
            chairperson = m
        else:
            others.append(m)
    rows = []
    if chairperson:
        rows.append([
            _fix( chairperson.get("name", "")),
            _fix( chairperson.get("position", "")),
            _fix( chairperson.get("address", "")),
            "", "", "",
        ])
    rows.append(["साक्षिहरू", "", "", "", "", ""])
    for m in others:
        rows.append([
            _fix( m.get("name", "")),
            _fix( m.get("position", "")),
            _fix( m.get("address", "")),
            "", "", "",
        ])
    return headers, rows


def _add_uc_members_table(doc: Document, val: list):
    headers, rows = _build_uc_members_data(val)
    if not headers or not rows:
        return
    num_rows = len(rows) + 1
    tbl = doc.add_table(rows=num_rows, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for ci, h in enumerate(headers):
        cell = tbl.cell(0, ci)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "006400")
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]._tr
        trPr = tr.get_or_add_trPr()
        h_elem = OxmlElement("w:trHeight")
        h_elem.set(qn("w:val"), "900")
        h_elem.set(qn("w:hRule"), "atLeast")
        trPr.append(h_elem)
        is_separator = ri >= 1 and row[0] == "साक्षिहरू"
        for ci, val_str in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            cell.text = val_str
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if is_separator:
                        r.font.bold = True
                if is_separator:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if is_separator:
                _set_cell_shading(cell, "f0f0f0")
    doc.add_paragraph()


def _add_uc_members_table_html(parts: list, val: list, var_name: str = ""):
    headers, rows = _build_uc_members_data(val)
    if not headers or not rows:
        return
    parts.append('<div class="table-preview"><table class="data" style="border-collapse:collapse;width:100%"><thead><tr>')
    for h in headers:
        parts.append(f'<th style="background:#006400;color:white;padding:8px;font-size:9pt;text-align:center;border:1px solid #006400">{_html_escape(h)}</th>')
    parts.append('</tr></thead><tbody>')
    for ri, row in enumerate(rows):
        is_separator = ri >= 1 and row[0] == "साक्षिहरू"
        bg = "#f0f0f0" if is_separator else "transparent"
        style = f'background:{bg};'
        parts.append(f'<tr style="{style}">')
        for ci, val_str in enumerate(row):
            align = "center" if is_separator and ci == 0 else "left"
            parts.append(f'<td style="padding:8px;font-size:9pt;border:1px solid #ddd;text-align:{align}">{_html_escape(val_str)}</td>')
        parts.append('</tr>')
    parts.append('</tbody></table></div>')

def _add_hh_records_table_html(parts: list, val: list, raw_data: dict = None):
    if not val or not isinstance(val, list):
        return
    hh = (raw_data or {}).get("households", {})
    total_hh = hh.get("total_households", len(val))
    total_pop = hh.get("total_population", 0)
    forest_occ = hh.get("forest_based_occupation", 0)
    parts.append(f'<h3 style="margin:16px 0 2px;font-size:14pt;font-weight:700;">{_html_escape(f"{format_devanagari(total_hh,0)} घरपरिवारको विस्तृत विवरण")}</h3>')
    parts.append(f'<p style="margin:0 0 8px;font-size:10pt;font-style:italic;">कुल जनसंख्या: {_html_escape(format_devanagari(total_pop,0))}  |  वनमा आश्रित: {_html_escape(format_devanagari(forest_occ,0))}</p>')
    keys = [k for k, _, _, _ in NP_HEADERS_HH_RECORDS]
    headers = [h for _, h, _, _ in NP_HEADERS_HH_RECORDS]
    col_types = [t for _, _, _, t in NP_HEADERS_HH_RECORDS]
    parts.append('<div class="table-preview"><table class="data" style="border-collapse:collapse;width:100%"><thead><tr>')
    for ci, np_header in enumerate(headers):
        display = _html_escape(np_header).replace("\n", "<br>")
        extra = ""
        if col_types[ci] == "number":
            extra = ' style="writing-mode:vertical-lr;text-orientation:mixed;white-space:nowrap;width:28px;height:80px;background:#006400;color:white;padding:4px 6px;font-size:9pt;text-align:center;border:1px solid #006400;"'
        elif col_types[ci] == "vertical":
            extra = ' style="writing-mode:vertical-lr;text-orientation:mixed;white-space:nowrap;width:28px;height:80px;background:#006400;color:white;padding:4px 6px;font-size:9pt;text-align:center;border:1px solid #006400;"'
        else:
            extra = ' style="background:#006400;color:white;padding:6px 8px;font-size:9pt;text-align:center;border:1px solid #006400;white-space:nowrap;"'
        parts.append(f'<th{extra}>{display}</th>')
    parts.append('</tr></thead><tbody>')
    for row in val:
        parts.append('<tr>')
        for ci, key in enumerate(keys):
            val_raw = row.get(key, "")
            if key == "जग्गा_एकाइ":
                val_str = "रो." if str(val_raw).strip().lower() == "ropani" else _html_escape(str(val_raw))
            elif key in ("घाँस_भारी", "दाउरा_भारी", "ओछ्यान_भारी"):
                val_str = _html_escape(format_devanagari(val_raw, 0))
            else:
                val_str = _html_escape(_fmt_value(val_raw, "hh_records"))
            align = "text-align:right" if col_types[ci] in ("number", "vertical") else "text-align:left"
            parts.append(f'<td style="padding:4px 6px;font-size:9pt;border:1px solid #ddd;{align};">{val_str}</td>')
        parts.append('</tr>')
    parts.append('</tbody></table></div>')


def _add_activity_plan_detail_table_html(parts: list, val: list):
    if not val or not isinstance(val, list):
        return
    keys = [eng_key for eng_key, _ in NP_HEADERS_ACTIVITY_PLAN if eng_key in val[0]]
    headers = [np_header for eng_key, np_header in NP_HEADERS_ACTIVITY_PLAN if eng_key in val[0]]
    parts.append('<div class="table-preview"><table class="data" style="border-collapse:collapse;width:100%"><thead><tr>')
    for np_header in headers:
        display = _html_escape(np_header).replace("\n", "<br>")
        parts.append(f'<th style="background:#006400;color:white;padding:6px 8px;font-size:9pt;text-align:center;border:1px solid #006400;white-space:nowrap;">{display}</th>')
    parts.append('</tr></thead><tbody>')
    for row in val:
        parts.append('<tr>')
        for key in keys:
            val_raw = row.get(key, "")
            if key == "total_budget":
                try:
                    val_raw = round(float(val_raw) / 1000)
                except (ValueError, TypeError):
                    pass
            parts.append(f'<td style="padding:6px 8px;font-size:9pt;border:1px solid #ddd;">{_html_escape(_fmt_value(val_raw, "ya_activity_plan_detail"))}</td>')
        parts.append('</tr>')
    parts.append('</tbody></table></div>')


def _add_static_table(doc: Document, node: TreeNode, raw_data: dict = None):
    data = node.static_table or {}
    caption = data.get("caption", "")
    columns = data.get("columns", [])
    rows = data.get("rows", [])
    merges = data.get("merges", []) or []
    if not columns or not rows:
        p = doc.add_paragraph()
        run = p.add_run("[Static table — no data]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(200, 0, 0)
        return
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_before = Pt(6)
        cap_para.paragraph_format.space_after = Pt(4)
        run = cap_para.add_run(_fix(caption))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = "Nirmala UI"
    tbl = doc.add_table(rows=1 + len(rows), cols=len(columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for ci, h in enumerate(columns):
        cell = tbl.cell(0, ci)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "006400")
    for m in merges:
        mr, mc = m["row"], m["col"]
        rs, cs = m["rowspan"], m["colspan"]
        if rs > 1 or cs > 1:
            master = tbl.cell(mr + 1, mc)
            target = tbl.cell(mr + rs, mc + cs - 1)
            merged = master.merge(target)
            cell_text = str(rows[mr][mc]) if rows and mr < len(rows) and mc < len(rows[mr]) and rows[mr][mc] is not None else ""
            if raw_data and cell_text.startswith("{{") and cell_text.endswith("}}"):
                cell_text = _resolve_var_text(cell_text, raw_data)
            merged.text = _fix(cell_text)
            for p in merged.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            is_slave = any(
                m["row"] <= ri - 1 < m["row"] + m["rowspan"]
                and m["col"] <= ci < m["col"] + m["colspan"]
                and not (ri - 1 == m["row"] and ci == m["col"])
                for m in merges
            )
            if is_slave:
                continue
            is_merged_master = any(
                ri - 1 == m["row"] and ci == m["col"] and (m["rowspan"] > 1 or m["colspan"] > 1)
                for m in merges
            )
            if is_merged_master:
                continue
            cell = tbl.cell(ri, ci)
            cell_text = str(val) if val is not None else ""
            if raw_data and cell_text.startswith("{{") and cell_text.endswith("}}"):
                cell_text = _resolve_var_text(cell_text, raw_data)
            cell.text = _fix(cell_text)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def _add_inline_tables(doc: Document, inline_tables: list, raw_data: dict = None):
    """Render multiple inline tables from a richtext node's inline_tables array."""
    for table_data in inline_tables:
        class _TmpNode:
            pass
        tmp = _TmpNode()
        tmp.static_table = table_data
        _add_static_table(doc, tmp, raw_data)


def _add_map(doc: Document, node: TreeNode, calculation_id: UUID, db: Session, calc_cache: dict = None):
    # Lazy imports (matplotlib + geoalchemy2 are heavy)
    from app.services.report.map_generator import generate_boundary_map
    from geoalchemy2.shape import to_shape
    from shapely.geometry import mapping

    if node.map_type:
        forest_name = (calc_cache or {}).get("forest_name", "CF")
        _add_map_standard(doc, node.map_type, calculation_id, db, forest_name)
        return

    calc = calc_cache.get("calculation") if calc_cache else None
    if not calc:
        calc = db.query(Calculation).filter(Calculation.id == calculation_id).first()
    if not calc:
        return

    forest_name = calc.forest_name or "CF"
    boundary_geojson = None
    blocks = []

    if calc.boundary_geom:
        try:
            shape = to_shape(calc.boundary_geom)
            boundary_geojson = mapping(shape)
        except Exception:
            pass

    cached_blocks = (calc_cache or {}).get("blocks")
    forest_blocks = cached_blocks if cached_blocks is not None else db.query(ForestBlock).filter(
        ForestBlock.calculation_id == calculation_id
    ).order_by(ForestBlock.index).all()

    for fb in forest_blocks:
        try:
            shape = to_shape(fb.geometry)
            geom = mapping(shape)
            centroid = {
                "lon": shape.centroid.x if hasattr(shape, 'centroid') else 0,
                "lat": shape.centroid.y if hasattr(shape, 'centroid') else 0,
            }
            blocks.append({
                "name": fb.name,
                "geometry": geom,
                "centroid": centroid,
                "area_hectares": fb.area_hectares,
            })
        except Exception:
            pass

    img_data = generate_boundary_map(boundary_geojson, forest_name, blocks if blocks else None)

    if img_data:
        try:
            if img_data.startswith("data:"):
                from app.utils.svg_to_png import add_svg_picture, _decode_data_uri
                raw_bytes, _ = _decode_data_uri(img_data)
                if len(raw_bytes) < 5000:
                    logger.warning(f"Boundary map too small ({len(raw_bytes)}B) — skipping blank image")
                    raise ValueError("Blank image")
                add_svg_picture(doc, raw_bytes, width_inches=5.5)
            else:
                if not os.path.exists(img_data) or os.path.getsize(img_data) < 5000:
                    logger.warning(f"Boundary map file too small or missing: {img_data}")
                    raise ValueError("Blank or missing image file")
                with open(img_data, "rb") as _f:
                    img_bytes = _f.read()
                from app.utils.svg_to_png import add_svg_picture
                add_svg_picture(doc, img_bytes, width_inches=5.5)
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(node.title_ne or "Forest Boundary Map")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()
            return
        except Exception:
            pass

    p = doc.add_paragraph()
    run = p.add_run(f"[Map: {node.title_ne}] - no boundary or block data available")
    run.font.italic = True
    run.font.color.rgb = RGBColor(200, 0, 0)


def _walk_tree(doc: Document, nodes: List[TreeNode], calculation_id: UUID,
               raw_data: Dict[str, Any], db: Session,
               table_cache: dict = None, calc_cache: dict = None):
    _wt_timers = {"text": 0.0, "chart": 0.0, "table": 0.0, "static_table": 0.0, "map": 0.0, "heading": 0.0, "children": 0.0, "total": 0.0}
    _wt_counts = {"text": 0, "chart": 0, "table": 0, "static_table": 0, "map": 0, "nodes": 0}
    _wt_slow_nodes = []
    _wt_start = time.time()
    for node in nodes:
        if node.hidden_in_export or node.deleted:
            continue

        has_content = node.content_type == "richtext" and node.content and node.content.strip()
        is_chart = node.content_type == "chart" and node.chart_type
        is_table = node.content_type == "table" and node.table_id
        is_static_table = node.content_type == "static_table"
        is_map = node.content_type == "map"
        has_children = any((not c.hidden_in_export and not c.deleted) for c in node.children)
        _wt_counts["nodes"] += 1
        _node_start = time.time()

        if has_content or is_chart or is_table or is_static_table or is_map or has_children:
            if node.type in ("section", "preamble", "appendix") or node.page_break_before:
                doc.add_page_break()
            _add_heading(doc, node)

        if has_content:
            _t0 = time.time()
            _add_text_content(doc, node.content, calculation_id, db, raw_data, table_cache)
            _wt_timers["text"] += time.time() - _t0
            _wt_counts["text"] += 1

            if node.inline_tables:
                _t0 = time.time()
                _add_inline_tables(doc, node.inline_tables, raw_data)
                _wt_timers["static_table"] += time.time() - _t0
                _wt_counts["static_table"] += 1

        if is_chart:
            _t0 = time.time()
            _add_chart(doc, node, raw_data, calculation_id)
            _wt_timers["chart"] += time.time() - _t0
            _wt_counts["chart"] += 1

        if is_table:
            _t0 = time.time()
            _add_table(doc, node, table_cache)
            _wt_timers["table"] += time.time() - _t0
            _wt_counts["table"] += 1

        if is_static_table:
            _t0 = time.time()
            _add_static_table(doc, node, raw_data)
            _wt_timers["static_table"] += time.time() - _t0
            _wt_counts["static_table"] += 1

        if is_map:
            _t0 = time.time()
            _add_map(doc, node, calculation_id, db, calc_cache)
            _wt_timers["map"] += time.time() - _t0
            _wt_counts["map"] += 1

        if has_children:
            _t0 = time.time()
            _walk_tree(doc, node.children, calculation_id, raw_data, db, table_cache, calc_cache)
            _wt_timers["children"] += time.time() - _t0

        if has_content or is_chart or is_table or is_static_table or is_map or has_children:
            doc.add_paragraph()

        _node_elapsed = time.time() - _node_start
        if _node_elapsed > 1.0:
            _wt_slow_nodes.append(f"{node.title_en or node.title_ne or '?'}={_node_elapsed:.1f}s")

    _wt_timers["total"] = time.time() - _wt_start
    if _wt_slow_nodes:
        logger.info("WT_SLOW: %s", " | ".join(_wt_slow_nodes))
    logger.info(
        "WT_PROFILE: nodes=%d text=%.2fs(%d) chart=%.2fs(%d) table=%.2fs(%d) static=%.2fs(%d) map=%.2fs(%d) heading=%.2fs children=%.2fs total=%.2fs",
        _wt_counts["nodes"],
        _wt_timers["text"], _wt_counts["text"],
        _wt_timers["chart"], _wt_counts["chart"],
        _wt_timers["table"], _wt_counts["table"],
        _wt_timers["static_table"], _wt_counts["static_table"],
        _wt_timers["map"], _wt_counts["map"],
        _wt_timers["heading"],
        _wt_timers["children"],
        _wt_timers["total"],
    )


# ═══════════════════════════════════════════════════════
# HTML Preview Functions
# ═══════════════════════════════════════════════════════

def _render_table_inline_replacement(table_id: str, table_cache: dict = None,
                                     calculation_id: UUID = None, db: Session = None) -> str:
    if table_id == "fieldinventory":
        return _render_fieldinventory_html(calculation_id, db)
    parts = []
    _add_table_inline_html(parts, table_id, table_cache)
    return "".join(parts) if parts else f'<div class="chart-placeholder">📋 {table_id}<br><small>Rendered as table in DOCX</small></div>'


def _render_fieldinventory_html(calculation_id: UUID, db: Session) -> str:
    """Render {{table:fieldinventory}} as HTML for preview."""
    if not calculation_id or not db:
        return '<div class="chart-placeholder">📋 क्षेत्र सर्वेक्षण मापन तथ्याङ्क<br><small>Rendered as table in DOCX</small></div>'

    fi_calc = db.query(FieldInventoryCalculation).filter(
        FieldInventoryCalculation.calculation_id == calculation_id
    ).first()
    if not fi_calc:
        return '<div class="chart-placeholder">📋 क्षेत्र सर्वेक्षण — no data</div>'

    measurements = (
        db.query(FieldInventoryMeasurement)
        .options(joinedload(FieldInventoryMeasurement.sample_plot))
        .join(FieldInventorySamplePlot)
        .filter(FieldInventorySamplePlot.field_inventory_calculation_id == fi_calc.id)
        .all()
    )
    if not measurements:
        return '<div class="chart-placeholder">📋 क्षेत्र सर्वेक्षण — no data</div>'

    measurements.sort(key=lambda m: (
        m.sample_plot.block_name,
        m.sample_plot.sample_plot_number,
        ST_TYPE_SORT.get(m.stand_type, 99),
        m.sn or 0,
    ))

    ST_TYPE_MAP_NP = {
        "Regeneration": "पुनरुत्पादन", "Sapling": "लाथ्रा",
        "Pole": "पोल", "Tree": "रूख",
    }

    h = _html_escape
    parts = ['<div class="table-preview">']
    parts.append('<h4 style="margin:12px 0 4px;font-size:14px;font-weight:700;">क्षेत्र सर्वेक्षण मापन तथ्याङ्क</h4>')
    parts.append('<table class="data"><thead><tr>')
    for np_h in ("ब्लक", "प्लट नं.", "प्रकार", "प्रजाति", "DBH (से.मि.)", "उचाइ (मि.)", "वर्ग", "गणना"):
        parts.append(f'<th>{np_h}</th>')
    parts.append('</tr></thead><tbody>')

    for m in measurements:
        sp = m.sample_plot
        if not sp:
            continue
        stand_type_np = ST_TYPE_MAP_NP.get(m.stand_type, m.stand_type or "—")
        species = (m.species_scientific or "")[:25]
        dbh_str = format_devanagari(m.dbh_cm, 1) if m.dbh_cm is not None else "—"
        ht_str = format_devanagari(m.height_m, 1) if m.height_m is not None else "—"
        cls_str = h(str(m.tree_class)) if m.tree_class else "—"
        cnt_str = format_devanagari(m.count, 0) if m.count else "1"
        parts.append('<tr>')
        parts.append(f'<td>{h(sp.block_name)}</td>')
        parts.append(f'<td>{format_devanagari(sp.sample_plot_number, 0)}</td>')
        parts.append(f'<td>{stand_type_np}</td>')
        parts.append(f'<td>{h(species)}</td>')
        parts.append(f'<td>{dbh_str}</td>')
        parts.append(f'<td>{ht_str}</td>')
        parts.append(f'<td>{cls_str}</td>')
        parts.append(f'<td>{cnt_str}</td>')
        parts.append('</tr>')
    parts.append('</tbody></table>')

    # Yield table
    seen_plots = set()
    yield_rows = []
    for m in measurements:
        sp = m.sample_plot
        if not sp:
            continue
        key = (sp.block_name, sp.sample_plot_number)
        if key not in seen_plots:
            seen_plots.add(key)
            yield_rows.append({
                "block_name": sp.block_name,
                "plot_no": format_devanagari(sp.sample_plot_number, 0),
                "firewood": format_devanagari(sp.firewood_kg_per_100sqm_per_year, 1) if sp.firewood_kg_per_100sqm_per_year is not None else "—",
                "grass": format_devanagari(sp.grass_kg_per_100sqm_per_year, 1) if sp.grass_kg_per_100sqm_per_year is not None else "—",
                "bedding": format_devanagari(sp.bedding_material_kg_per_100sqm_per_year, 1) if sp.bedding_material_kg_per_100sqm_per_year is not None else "—",
            })

    if yield_rows:
        parts.append('<h4 style="margin:12px 0 4px;font-size:14px;font-weight:700;">नमुना प्लटमा उपलब्ध स्रोत परिणाम</h4>')
        parts.append('<table class="data"><thead><tr>')
        for yh in ("ब्लक", "प्लट नं.", "दाउरा (के.जी.)", "घाँस (के.जी.)", "सोतर (के.जी.)"):
            parts.append(f'<th>{yh}</th>')
        parts.append('</tr></thead><tbody>')
        for yr in yield_rows:
            parts.append('<tr>')
            parts.append(f'<td>{h(yr["block_name"])}</td>')
            parts.append(f'<td>{yr["plot_no"]}</td>')
            parts.append(f'<td>{yr["firewood"]}</td>')
            parts.append(f'<td>{yr["grass"]}</td>')
            parts.append(f'<td>{yr["bedding"]}</td>')
            parts.append('</tr>')
        parts.append('</tbody></table>')

    parts.append('</div>')
    return "\n".join(parts)


def _render_chart_html(chart_type: str, raw_data: dict, calculation_id: UUID = None, forest_name: str = "") -> str:
    """Render a chart as base64-embedded PNG for HTML preview."""
    if calculation_id:
        cached = _chart_cache_get(calculation_id, chart_type)
        if cached:
            import base64
            b64 = base64.b64encode(cached.getvalue()).decode("utf-8")
            return f'<div style="margin:12px 0;text-align:center;"><img src="data:image/svg+xml;base64,{b64}" style="max-width:100%;border:1px solid #ddd;border-radius:4px;" alt="{_html_escape(chart_type)}"><div style="font-size:9pt;color:#666;margin-top:4px;">{chart_type.replace("_", " ").title()}</div></div>'

    if not forest_name:
        forest_name = raw_data.get("basic_info", {}).get("forest_name", "")

    img_data = _generate_chart(chart_type, raw_data, forest_name)

    if img_data:
        try:
            if img_data.startswith("data:"):
                if calculation_id:
                    from app.utils.svg_to_png import _decode_data_uri
                    raw_bytes, _ = _decode_data_uri(img_data)
                    _chart_cache_set(calculation_id, chart_type, BytesIO(raw_bytes))
                return f'<div style="margin:12px 0;text-align:center;"><img src="{img_data}" style="max-width:100%;border:1px solid #ddd;border-radius:4px;" alt="{_html_escape(chart_type)}"><div style="font-size:9pt;color:#666;margin-top:4px;">{chart_type.replace("_", " ").title()}</div></div>'
            else:
                with open(img_data, "rb") as _f:
                    img_bytes = _f.read()
                if calculation_id:
                    _chart_cache_set(calculation_id, chart_type, BytesIO(img_bytes))
                import base64
                b64 = base64.b64encode(img_bytes).decode("utf-8")
                return f'<div style="margin:12px 0;text-align:center;"><img src="data:image/svg+xml;base64,{b64}" style="max-width:100%;border:1px solid #ddd;border-radius:4px;" alt="{_html_escape(chart_type)}"><div style="font-size:9pt;color:#666;margin-top:4px;">{chart_type.replace("_", " ").title()}</div></div>'
        except Exception:
            pass

    return f'<div class="chart-placeholder">📊 {chart_type}<br><small>Chart data not available</small></div>'


def _render_map_html(calculation_id: UUID, db: Session, layer_name: str, forest_name: str = "CF") -> str:
    """Render a map layer as base64-embedded PNG for HTML preview."""
    if not calculation_id or not db:
        return f'<div class="chart-placeholder">🗺️ {layer_name}<br><small>No data</small></div>'
    try:
        from app.services.management_plan_docx.plan_map_service import generate_standard_map
        buf = generate_standard_map(db, calculation_id, layer_name, forest_name=forest_name, use_cache=True)
        if buf and buf.getbuffer().nbytes > 5000:
            import base64
            buf.seek(0)
            b64 = base64.b64encode(buf.read()).decode("utf-8")
            return f'<div style="margin:12px 0;text-align:center;"><img src="data:image/png;base64,{b64}" style="max-width:100%;border:1px solid #ddd;border-radius:4px;" alt="{_html_escape(layer_name)}"></div>'
    except Exception as e:
        logger.warning(f"Map preview failed for {layer_name}: {e}")
    return f'<div class="chart-placeholder">🗺️ {layer_name}<br><small>Maps are generated in DOCX/PDF export. No preview available.</small></div>'


def _walk_tree_html(nodes: List[TreeNode], calculation_id: UUID,
                    raw_data: Dict[str, Any], db: Session,
                    table_cache: dict = None) -> str:
    parts = []
    forest_name = ""
    if raw_data:
        bi = raw_data.get("basic_info", {})
        if isinstance(bi, dict):
            forest_name = bi.get("forest_name", "")
    for node in nodes:
        if node.hidden_in_export or node.deleted:
            continue
        has_content = node.content_type == "richtext" and node.content and node.content.strip()
        is_chart = node.content_type == "chart" and node.chart_type
        is_table = node.content_type == "table" and node.table_id
        is_static_table = node.content_type == "static_table"
        is_map = node.content_type == "map"
        has_children = any((not c.hidden_in_export and not c.deleted) for c in node.children)
        if not (has_content or is_chart or is_table or is_static_table or is_map or has_children):
            continue

        num = f"{node.number}. " if node.number else ""
        tag = "h2" if node.type in ("subsection",) else "h1"
        page_cls = " page-break" if node.page_break_before else ""
        parts.append(f'<div class="section{page_cls}" id="{node.id}">')
        if node.page_break_before:
            parts.append('<hr class="pb-marker">')
        parts.append(f'<{tag}>{num}{node.title_ne}</{tag}>')
        if node.title_en and node.title_en != node.title_ne:
            parts.append(f'<p style="color:#666;font-style:italic;font-size:0.85em;margin:0 0 8px;">{_html_escape(node.title_en)}</p>')

        if has_content:
            content = _fix(node.content)
            escaped = _html_escape(content)
            escaped = re.sub(
                r'\{\{chart:(\w+)\}\}',
                lambda m: _render_chart_html(m.group(1), raw_data, calculation_id, forest_name),
                escaped
            )
            escaped = re.sub(
                r'\{\{map:(\w+)\}\}',
                lambda m: _render_map_html(calculation_id, db, m.group(1), forest_name),
                escaped
            )
            escaped = re.sub(
                r'\{\{table:(\w+)\}\}',
                lambda m: _render_table_inline_replacement(m.group(1), table_cache, calculation_id, db),
                escaped
            )
            escaped = _render_html_list_vars(escaped, raw_data)
            escaped = re.sub(
                r'\{\{section:(\w+):full\}\}',
                lambda m: _render_section_full_html(m.group(1), raw_data, calculation_id),
                escaped
            )
            parts.append(f'<div class="section-content">{escaped}</div>')

            if node.inline_tables:
                for table_data in node.inline_tables:
                    _add_inline_table_html(parts, table_data, raw_data)

        if is_chart:
            parts.append(_render_chart_html(node.chart_type, raw_data, calculation_id, forest_name))

        if is_map:
            layer_name = node.map_type or "boundary"
            parts.append(_render_map_html(calculation_id, db, layer_name, forest_name))

        if is_table:
            _add_table_html(parts, node, table_cache)

        if is_static_table:
            _add_static_table_html(parts, node, raw_data)

        if has_children:
            parts.append(_walk_tree_html(node.children, calculation_id, raw_data, db, table_cache))

        parts.append('</div>')
    return "\n".join(parts)


def _add_table_html(parts: List[str], node: TreeNode, table_cache: dict = None):
    table_id = node.table_id
    if not table_id:
        return
    vdef = get_variable(f"table:{table_id}")
    title_ne = vdef.label_ne if vdef else table_id.replace("_", " ").title()
    parts.append(f'<h3 style="margin:16px 0 2px;font-size:14pt;font-weight:700;">{_html_escape(title_ne)}</h3>')
    table_data = (table_cache or {}).get(table_id)
    if not table_data or not table_data.rows:
        parts.append(f'<div class="chart-placeholder">📋 {table_id} — no data</div>')
        return
    rows = table_data.rows
    if not rows:
        return
    headers = list(rows[0].keys()) if isinstance(rows[0], dict) else []
    np_headers = NP_HEADERS_BIODIVERSITY.get(table_id, [])
    np_map = {eng: np for eng, np in np_headers}
    parts.append('<div class="table-preview"><table class="data"><thead><tr>')
    for h in headers:
        display = np_map.get(h, h.replace("_", " ").title())
        parts.append(f'<th>{_html_escape(display)}</th>')
    parts.append('</tr></thead><tbody>')
    for row in rows:
        parts.append('<tr>')
        for h in headers:
            val = format_devanagari(row.get(h, ""))
            parts.append(f'<td>{_html_escape(val)}</td>')
        parts.append('</tr>')
    parts.append('</tbody></table></div>')
    if node.title_ne:
        parts.append(f'<p style="text-align:center;font-style:italic;font-size:9pt;color:#666;margin:4px 0 12px;">{_html_escape(node.title_ne)}</p>')


def _add_table_inline_html(parts: list, table_id: str, table_cache: dict = None):
    from app.services.operational_plan.variable_registry import TABLE_ID_ALIAS
    table_id = TABLE_ID_ALIAS.get(table_id, table_id)
    vdef = get_variable(f"table:{table_id}")
    title_ne = vdef.label_ne if vdef else table_id.replace("_", " ").title()
    parts.append(f'<h3 style="margin:16px 0 2px;font-size:14pt;font-weight:700;">{_html_escape(title_ne)}</h3>')
    table_data = (table_cache or {}).get(table_id)
    if not table_data or not table_data.rows:
        parts.append(f'<div class="chart-placeholder">📋 {table_id} — no data</div>')
        return
    rows = table_data.rows
    headers = list(rows[0].keys()) if rows and isinstance(rows[0], dict) else []
    np_headers = NP_HEADERS_BIODIVERSITY.get(table_id, [])
    np_map = {eng: np for eng, np in np_headers}
    parts.append('<div class="table-preview"><table class="data"><thead><tr>')
    for h in headers:
        display = np_map.get(h, h.replace("_", " ").title())
        parts.append(f'<th>{_html_escape(display)}</th>')
    parts.append('</tr></thead><tbody>')
    for row in rows:
        parts.append('<tr>')
        for h in headers:
            val = format_devanagari(row.get(h, ""))
            parts.append(f'<td>{_html_escape(val)}</td>')
        parts.append('</tr>')
    parts.append('</tbody></table></div>')


def _add_static_table_html(parts: List[str], node: TreeNode, raw_data: dict = None):
    data = node.static_table or {}
    caption = data.get("caption", "")
    columns = data.get("columns", [])
    rows = data.get("rows", [])
    merges = data.get("merges", []) or []
    if not columns or not rows:
        parts.append('<div class="chart-placeholder">📋 Static table — no data</div>')
        return
    if caption:
        parts.append(f'<p style="text-align:center;font-weight:bold;font-size:10pt;margin:8px 0 4px;">{_html_escape(caption)}</p>')

    visible = {}
    for m in merges:
        mr, mc = m["row"], m["col"]
        rs, cs = m["rowspan"], m["colspan"]
        for r in range(mr, mr + rs):
            for c in range(mc, mc + cs):
                visible[(r, c)] = None if (r != mr or c != mc) else (rs, cs)

    parts.append('<div class="table-preview"><table class="data"><thead><tr>')
    for ci, h in enumerate(columns):
        key = (0, ci)
        if key in visible:
            if visible[key] is None:
                continue
            rs, cs = visible[key]
            parts.append(f'<th rowspan="{rs}" colspan="{cs}">{_html_escape(h)}</th>')
        else:
            parts.append(f'<th>{_html_escape(h)}</th>')
    parts.append('</tr></thead><tbody>')

    for ri, row in enumerate(rows):
        parts.append('<tr>')
        for ci, val in enumerate(row):
            key = (ri, ci)
            if key in visible:
                if visible[key] is None:
                    continue
                rs, cs = visible[key]
            else:
                rs, cs = 1, 1
            cell_text = str(val) if val is not None else ""
            if raw_data and cell_text.startswith("{{") and cell_text.endswith("}}"):
                cell_text = _resolve_var_text(cell_text, raw_data)
            attrs = f' rowspan="{rs}" colspan="{cs}"' if (rs > 1 or cs > 1) else ''
            parts.append(f'<td{attrs}>{_html_escape(cell_text)}</td>')
        parts.append('</tr>')
    parts.append('</tbody></table></div>')


def _add_inline_table_html(parts: List[str], table_data: dict, raw_data: dict = None):
    """Render a single inline table dict as HTML."""
    caption = table_data.get("caption", "")
    columns = table_data.get("columns", [])
    rows = table_data.get("rows", [])
    merges = table_data.get("merges", []) or []
    if not columns or not rows:
        return
    if caption:
        parts.append(f'<p style="text-align:center;font-weight:bold;font-size:10pt;margin:8px 0 4px;">{_html_escape(caption)}</p>')
    class _TmpNode:
        pass
    tmp = _TmpNode()
    tmp.static_table = table_data
    _add_static_table_html(parts, tmp, raw_data)


_SECTION_TITLES = {
    "forest_summary": ("वन सारांश", "Forest Summary"),
    "slope_analysis": ("भिरालो विश्लेषण", "Slope Analysis"),
    "elevation_profile": ("उचाइ विवरण", "Elevation Profile"),
    "aspect_analysis": ("दिशा विश्लेषण", "Aspect Analysis"),
    "forest_health": ("वन स्वास्थ्य", "Forest Health"),
    "forest_type": ("वन प्रकार", "Forest Type"),
    "species_potential": ("सम्भावित प्रजातिहरू", "Potential Species"),
    "actual_species": ("वास्तविक प्रजातिहरू", "Actual Species"),
    "biodiversity": ("जैविक विविधता", "Biodiversity"),
    "canopy_structure": ("वन मुकुट", "Canopy Structure"),
    "biomass_carbon": ("जैविक पदार्थ तथा कार्बन", "Biomass & Carbon"),
    "climate_conditions": ("मौसम अवस्था", "Climate Conditions"),
    "land_cover": ("भू-आवरण", "Land Cover"),
    "forest_loss": ("वन क्षति", "Forest Loss"),
    "fire_loss": ("आगलागी क्षति", "Fire Loss"),
    "forest_quality": ("वन गुणस्तर (नासा)", "Forest Quality"),
    "soil_analysis": ("माटो विश्लेषण", "Soil Analysis"),
    "location_context": ("स्थान तथा सन्दर्भ", "Location & Context"),
    "species_distribution": ("प्रजाति वितरण", "Species Distribution"),
    "accessible_forest": ("पहुँचयोग्य वन क्षेत्र", "Accessible Forest"),
}

_SECTION_CHARTS = {
    "slope_analysis": "slope_bar",
    "aspect_analysis": "aspect_rose",
    "forest_health": "forest_health_pie",
    "forest_type": "forest_type_pie",
    "actual_species": "species_composition_pie_fi",
    "biodiversity": "species_composition_pie_fi",
    "canopy_structure": "canopy_bar",
    "land_cover": "landcover_pie",
    "forest_loss": None,
    "fire_loss": None,
    "forest_quality": "nasa_forest_2020_pie",
    "soil_analysis": "soil_bar",
}


def _render_section_full_html(section_name: str, raw_data: dict, calculation_id) -> str:
    key = f"section:{section_name}"
    sections = raw_data.get("section_generators", {})
    narrative = sections.get(key, "")
    if not narrative:
        return ""
    title_np, title_en = _SECTION_TITLES.get(section_name, (section_name, section_name))
    parts_html = [f'<div class="section-full" id="section-full-{section_name}">']
    parts_html.append(f'<h3>{_html_escape(title_np)}</h3>')
    parts_html.append(f'<p><small><em>{_html_escape(title_en)}</em></small></p>')
    parts_html.append(f'<div class="section-full-narrative"><p>{_html_escape(narrative)}</p></div>')
    chart_type = _SECTION_CHARTS.get(section_name)
    if chart_type:
        parts_html.append(_render_chart_html(chart_type, raw_data, calculation_id))
    parts_html.append('</div>')
    return "\n".join(parts_html)


def _html_escape(text: str) -> str:
    return (text.replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


def _render_html_list_vars(text: str, raw_data: dict) -> str:
    def _replace_var(m):
        var_name = m.group(1)
        if var_name.startswith("chart:") or var_name.startswith("map:") or var_name.startswith("table:") or var_name.endswith(":full"):
            return m.group(0)
        var_val = _resolve_var_from_raw(var_name, raw_data)
        if var_val is None:
            return ""
        if isinstance(var_val, list):
            if all(isinstance(v, str) for v in var_val):
                items = "".join(f"<li>{_html_escape(v)}</li>" for v in var_val if v)
                return f"<ul>{items}</ul>" if items else ""
            if all(isinstance(v, dict) for v in var_val):
                if var_name == "uc_members":
                    parts = []
                    _add_uc_members_table_html(parts, var_val, var_name)
                    return "".join(parts)
                if var_name == "ya_activity_plan_detail":
                    parts = []
                    _add_activity_plan_detail_table_html(parts, var_val)
                    return "".join(parts)
                if var_name == "hh_records":
                    parts = []
                    _add_hh_records_table_html(parts, var_val, raw_data)
                    return "".join(parts)
                title_html = ""
                vdef = get_variable(var_name)
                if vdef and vdef.label_ne:
                    title_html = f'<h4 style="margin:12px 0 4px;font-size:14px;font-weight:700;">{_html_escape(vdef.label_ne)}</h4>'
                headers = list(var_val[0].keys())
                header_row = "".join(f"<th>{_html_escape(h.replace('_', ' ').title())}</th>" for h in headers)
                data_rows = ""
                for row in var_val:
                    cells = "".join(f"<td>{_html_escape(_fmt_value(row.get(h, ''), var_name))}</td>" for h in headers)
                    data_rows += f"<tr>{cells}</tr>"
                return f'{title_html}<div class="table-preview"><table class="data"><thead><tr>{header_row}</tr></thead><tbody>{data_rows}</tbody></table></div>'
            items = "".join(f"<li>{_html_escape(str(v))}</li>" for v in var_val if v)
            return f"<ul>{items}</ul>" if items else m.group(0)
        if isinstance(var_val, dict):
            items = "".join(f"<li><b>{_html_escape(k)}</b>: {_html_escape(_fmt_value(v, var_name))}</li>" for k, v in var_val.items() if v)
            return f"<ul>{items}</ul>" if items else m.group(0)
        return _html_escape(_fmt_value(var_val, var_name))
    return re.sub(r"\{\{(\w+(?::\w+)*)\}\}", _replace_var, text)


# ═══════════════════════════════════════════════════════
# DOCX Builder
# ═══════════════════════════════════════════════════════

class _RawDataTableProxy:
    """Wraps a dict as a fake OPTableData object for fallback table_cache entries."""
    def __init__(self, table_id: str, rows: list):
        self.table_id = table_id
        self.rows = rows
        self.auto_populated = True

def _build_table_cache(calculation_id: UUID, db: Session, raw_data: dict = None) -> dict:
    all_tables = db.query(OPTableData).filter(
        OPTableData.calculation_id == calculation_id
    ).all()
    cache = {t.table_id: t for t in all_tables}
    # Supplement missing tables from data_collector raw data
    if raw_data:
        for t_id, rows_fn in _BUILD_FROM_RAW.items():
            if t_id not in cache:
                rows = rows_fn(raw_data)
                if rows:
                    cache[t_id] = _RawDataTableProxy(t_id, rows)
    # Override demand_supply with raw data fallback to ensure chart/table consistency
    # (auto-populated OPTableData may have stale/incomplete data from old code)
    if raw_data and "demand_supply" in _BUILD_FROM_RAW:
        rows = _BUILD_FROM_RAW["demand_supply"](raw_data)
        if rows:
            cache["demand_supply"] = _RawDataTableProxy("demand_supply", rows)
    return cache


_DS_PRODUCT_LABELS = {
    "firewood_bhari": "दाउरा भारी",
    "grass_bhari": "घाँस भारी",
    "bedding_bhari": "सोतर भारी",
    "timber_cft": "काठ क्यू.फि.",
    "poles_count": "खाँवा संख्या",
}

def _build_demand_supply_rows(raw_data: dict) -> list:
    ds = raw_data.get("demand_supply", {})
    if not ds.get("demand"):
        return []
    products = ["firewood_bhari", "grass_bhari", "bedding_bhari", "timber_cft", "poles_count"]
    rows = []
    for k in products:
        cf_reg = ds.get("supply_cf_regular", {})
        cf_aah = ds.get("supply_cf_aah", {})
        row = {
            "product": _DS_PRODUCT_LABELS.get(k, k),
            "demand": ds.get("demand", {}).get(k, 0) or 0,
            "cf_regular": cf_reg[k] if k in cf_reg else "-",
            "cf_aah": cf_aah[k] if k in cf_aah else "-",
            "private": ds.get("supply_private", {}).get(k, 0) or 0,
            "total_supply": ds.get("total_supply", {}).get(k, 0) or 0,
        }
        deficit = ds.get("deficit", {}).get(k, 0) or 0
        if isinstance(deficit, (int, float)):
            sign = "बचत" if deficit >= 0 else "कमी"
            row["deficit"] = f"{sign} {abs(deficit):.2f}"
        else:
            row["deficit"] = str(deficit) if deficit else "-"
        rows.append(row)
    return rows


def _build_biodiversity_table_rows(raw_data: dict) -> list:
    bio = raw_data.get("biodiversity", {})
    if not bio.get("available"):
        return []
    rows = []
    idx = 0
    for rec in bio.get("vegetation", []):
        idx += 1
        rows.append({
            "sn": idx,
            "name": rec.get("name", ""),
            "scientific_name": rec.get("scientific_name", ""),
            "type": "वनस्पति",
            "sub_category": rec.get("sub_category", ""),
            "iucn_status": rec.get("iucn_status", ""),
            "is_protected": "हो" if rec.get("is_protected") else "होइन",
            "is_invasive": "हो" if rec.get("is_invasive") else "होइन",
        })
    for rec in bio.get("animals", []):
        idx += 1
        rows.append({
            "sn": idx,
            "name": rec.get("name", ""),
            "scientific_name": rec.get("scientific_name", ""),
            "type": "जनावर",
            "sub_category": rec.get("sub_category", ""),
            "iucn_status": rec.get("iucn_status", ""),
            "is_protected": "हो" if rec.get("is_protected") else "होइन",
            "is_invasive": "हो" if rec.get("is_invasive") else "होइन",
        })
    return rows


def _build_iucn_status_rows(raw_data: dict) -> list:
    bio = raw_data.get("biodiversity", {})
    if not bio.get("available"):
        return []
    iucn_map = {"CR": "संकटग्रस्त", "EN": "लोपोन्मुख", "VU": "असुरक्षित",
                "NT": "नजिकै खतरा", "LC": "कम चासो", "DD": "अपर्याप्त"}
    iucn_order = ["CR", "EN", "VU", "NT", "LC", "DD"]
    breakdown = bio.get("iucn_breakdown", {})
    rows = []
    for code in iucn_order:
        cnt = breakdown.get(code, 0)
        if cnt:
            rows.append({
                "iucn_code": code,
                "nepali_label": iucn_map.get(code, code),
                "count": cnt,
            })
    return rows


def _build_protected_species_rows(raw_data: dict) -> list:
    bio = raw_data.get("biodiversity", {})
    if not bio.get("available"):
        return []
    rows = []
    idx = 0
    for rec in bio.get("vegetation", []):
        if rec.get("is_protected"):
            idx += 1
            rows.append({
                "sn": idx,
                "name": rec.get("name", ""),
                "scientific_name": rec.get("scientific_name", ""),
                "sub_category": rec.get("sub_category", ""),
                "iucn_status": rec.get("iucn_status", ""),
            })
    for rec in bio.get("animals", []):
        if rec.get("is_protected"):
            idx += 1
            rows.append({
                "sn": idx,
                "name": rec.get("name", ""),
                "scientific_name": rec.get("scientific_name", ""),
                "sub_category": rec.get("sub_category", ""),
                "iucn_status": rec.get("iucn_status", ""),
            })
    return rows


def _build_invasive_species_rows(raw_data: dict) -> list:
    bio = raw_data.get("biodiversity", {})
    if not bio.get("available"):
        return []
    rows = []
    idx = 0
    for rec in bio.get("vegetation", []):
        if rec.get("is_invasive"):
            idx += 1
            rows.append({
                "sn": idx,
                "name": rec.get("name", ""),
                "scientific_name": rec.get("scientific_name", ""),
                "sub_category": rec.get("sub_category", ""),
                "iucn_status": rec.get("iucn_status", ""),
            })
    for rec in bio.get("animals", []):
        if rec.get("is_invasive"):
            idx += 1
            rows.append({
                "sn": idx,
                "name": rec.get("name", ""),
                "scientific_name": rec.get("scientific_name", ""),
                "sub_category": rec.get("sub_category", ""),
                "iucn_status": rec.get("iucn_status", ""),
            })
    return rows


def _build_vegetation_species_rows(raw_data: dict) -> list:
    bio = raw_data.get("biodiversity", {})
    if not bio.get("available"):
        return []
    rows = []
    for idx, rec in enumerate(bio.get("vegetation", []), 1):
        rows.append({
            "sn": idx,
            "name": rec.get("name", ""),
            "scientific_name": rec.get("scientific_name", ""),
            "sub_category": rec.get("sub_category", ""),
            "iucn_status": rec.get("iucn_status", ""),
            "is_protected": "हो" if rec.get("is_protected") else "होइन",
            "is_invasive": "हो" if rec.get("is_invasive") else "होइन",
            "primary_use": rec.get("primary_use", ""),
        })
    return rows


def _build_animal_species_rows(raw_data: dict) -> list:
    bio = raw_data.get("biodiversity", {})
    if not bio.get("available"):
        return []
    rows = []
    for idx, rec in enumerate(bio.get("animals", []), 1):
        rows.append({
            "sn": idx,
            "name": rec.get("name", ""),
            "scientific_name": rec.get("scientific_name", ""),
            "sub_category": rec.get("sub_category", ""),
            "iucn_status": rec.get("iucn_status", ""),
            "is_protected": "हो" if rec.get("is_protected") else "होइन",
            "is_invasive": "हो" if rec.get("is_invasive") else "होइन",
            "primary_use": rec.get("primary_use", ""),
        })
    return rows


# ═══════════════════════════════════════════════════════════════════
# Tree Mapping Analysis builder functions
# ═══════════════════════════════════════════════════════════════════

def _build_sm_hierarchy_summary(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_hierarchy_summary", [])

def _build_sm_species_by_hierarchy(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_species_by_hierarchy", [])

def _build_sm_species_diversity(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_species_diversity", [])

def _build_sm_dbh_by_hierarchy(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_dbh_by_hierarchy", [])

def _build_sm_dbh_species_by_hierarchy(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_dbh_species_by_hierarchy", [])

def _build_sm_stand_type_by_hierarchy(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_stand_type_by_hierarchy", [])

def _build_sm_carbon_by_hierarchy(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_carbon_by_hierarchy", [])

def _build_sm_volume_by_hierarchy(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_volume_by_hierarchy", [])

def _build_sm_top_species_by_volume(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_top_species_by_volume", [])

def _build_sm_mother_tree_by_hierarchy(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_mother_tree_by_hierarchy", [])

def _build_sm_mother_tree_by_species(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_mother_tree_by_species", [])

def _build_sm_felling_tree_by_species(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_felling_tree_by_species", [])

def _build_sm_species_hier_remark(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_species_hier_remark", [])

def _build_sm_dbh_hier_remark(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_dbh_hier_remark", [])

def _build_sm_felling_dbh_analysis(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_felling_dbh_analysis", [])

def _build_sm_felling_species_analysis(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_felling_species_analysis", [])

def _build_sm_mf_hierarchy_legend(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_mf_hierarchy_legend", [])

def _build_sm_stand_type_legend(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_stand_type_legend", [])

def _build_sm_carbon_legend(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_carbon_legend", [])

def _build_sm_volume_legend(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_volume_legend", [])

def _build_sm_mf_species_legend(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_mf_species_legend", [])

def _build_sm_felling_species_legend(raw_data: dict) -> list:
    sm = raw_data.get("tree_mapping_analysis", {})
    return sm.get("sm_felling_species_legend", [])


_BUILD_FROM_RAW = {
    "demand_supply": _build_demand_supply_rows,
    "table_20": _build_biodiversity_table_rows,
    "table_33": _build_iucn_status_rows,
    "table_34": _build_protected_species_rows,
    "table_35": _build_invasive_species_rows,
    "table_36": _build_vegetation_species_rows,
    "table_37": _build_animal_species_rows,
    "sm_hierarchy_summary": _build_sm_hierarchy_summary,
    "sm_species_by_hierarchy": _build_sm_species_by_hierarchy,
    "sm_species_diversity": _build_sm_species_diversity,
    "sm_dbh_by_hierarchy": _build_sm_dbh_by_hierarchy,
    "sm_dbh_species_by_hierarchy": _build_sm_dbh_species_by_hierarchy,
    "sm_stand_type_by_hierarchy": _build_sm_stand_type_by_hierarchy,
    "sm_carbon_by_hierarchy": _build_sm_carbon_by_hierarchy,
    "sm_volume_by_hierarchy": _build_sm_volume_by_hierarchy,
    "sm_top_species_by_volume": _build_sm_top_species_by_volume,
    "sm_mother_tree_by_hierarchy": _build_sm_mother_tree_by_hierarchy,
    "sm_mother_tree_by_species": _build_sm_mother_tree_by_species,
    "sm_felling_tree_by_species": _build_sm_felling_tree_by_species,
    "sm_species_hier_remark": _build_sm_species_hier_remark,
    "sm_dbh_hier_remark": _build_sm_dbh_hier_remark,
    "sm_felling_dbh_analysis": _build_sm_felling_dbh_analysis,
    "sm_felling_species_analysis": _build_sm_felling_species_analysis,
    "sm_mf_hierarchy_legend": _build_sm_mf_hierarchy_legend,
    "sm_stand_type_legend": _build_sm_stand_type_legend,
    "sm_carbon_legend": _build_sm_carbon_legend,
    "sm_volume_legend": _build_sm_volume_legend,
    "sm_mf_species_legend": _build_sm_mf_species_legend,
    "sm_felling_species_legend": _build_sm_felling_species_legend,
}


def _build_calc_cache(calculation_id: UUID, db: Session) -> dict:
    calc = db.query(Calculation).filter(Calculation.id == calculation_id).first()
    blocks = db.query(ForestBlock).filter(
        ForestBlock.calculation_id == calculation_id
    ).order_by(ForestBlock.index).all()
    return {
        "calculation": calc,
        "blocks": blocks,
        "forest_name": calc.forest_name if calc else "CF",
    }


def build_op_document(
    plan: Dict[str, Any],
    tree: List[TreeNode],
    resolver: VariableResolver,
    calculation_id: UUID,
    db: Session,
) -> BytesIO:
    t0 = time.time()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Nirmala UI"
    style.font.size = Pt(11)

    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Nirmala UI")

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    metadata = plan.get("plan_metadata", {})
    _add_cover_page(doc, plan, metadata)
    _add_toc_field(doc)
    t1 = time.time()

    raw_data = resolver.get_raw_data()
    raw_data["user_inputs"] = metadata.get("user_inputs", {})
    table_cache = _build_table_cache(calculation_id, db, raw_data)
    calc_cache = _build_calc_cache(calculation_id, db)
    t2 = time.time()

    _walk_tree(doc, tree, calculation_id, raw_data, db, table_cache, calc_cache)
    t3 = time.time()

    custom_notes = metadata.get("custom_notes")
    if custom_notes:
        doc.add_page_break()
        doc.add_heading("Custom Notes", level=1)
        for line in custom_notes.strip().split("\n"):
            doc.add_paragraph(line.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    t4 = time.time()

    logger.info(
        "OP_BUILD: setup=%.2fs cache=%.2fs walk_tree=%.2fs save=%.2fs total=%.2fs",
        t1 - t0, t2 - t1, t3 - t2, t4 - t3, t4 - t0,
    )
    return buffer
