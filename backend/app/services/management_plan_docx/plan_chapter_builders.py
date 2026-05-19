import io
from typing import Dict, Any, Optional, List
from uuid import UUID
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches, Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from .plan_map_service import generate_standard_map, LAYER_LABELS
from .plan_raster_descriptions import describe_layer
from .plan_narrative import get_chapter_narrative
from ..field_inventory_mgmt_charts import (
    chart_species_composition,
    chart_block_comparison,
    chart_annual_harvest,
    chart_forest_condition,
    chart_dbh_class_volume,
    chart_carbon_stock,
    chart_growth_rate,
    chart_stand_structure,
    chart_productivity,
)
import base64
from ..report.chart_generator import (
    generate_slope_pie as _slope_pie_b64,
    generate_landcover_pie as _lc_pie_b64,
    generate_aspect_rose as _aspect_rose_b64,
    generate_forest_type_pie as _ft_pie_b64,
)


def _b64_to_buf(fn_result: str) -> io.BytesIO:
    """Convert a base64 data URI from chart_generator to BytesIO."""
    if not fn_result:
        return None
    if fn_result.startswith('data:'):
        data = fn_result.split(',')[1]
        return io.BytesIO(base64.b64decode(data))
    return io.BytesIO(fn_result.encode())


def generate_slope_pie(slope_pcts, forest_name=""):
    return _b64_to_buf(_slope_pie_b64(slope_pcts, forest_name=forest_name))


def generate_landcover_pie(lc_pcts, forest_name=""):
    return _b64_to_buf(_lc_pie_b64(lc_pcts, forest_name=forest_name))


def generate_aspect_rose(asp_pcts, forest_name=""):
    return _b64_to_buf(_aspect_rose_b64(asp_pcts, forest_name=forest_name))


def generate_forest_type_pie(ft_pcts, forest_name=""):
    return _b64_to_buf(_ft_pie_b64(ft_pcts, forest_name=forest_name))

COLOR_GREEN = RGBColor(0, 100, 0)
COLOR_DARK = RGBColor(33, 33, 33)
COLOR_GRAY = RGBColor(128, 128, 128)
COLOR_WHITE = RGBColor(255, 255, 255)
IMG_WIDTH_MM = 150


def _set_cell_shading(cell, color_hex: str):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def add_section_heading(doc: Document, num: str, title_ne: str, title_en: str = ""):
    heading = doc.add_heading(f'{num}. {title_ne}', level=1)
    for run in heading.runs:
        run.font.color.rgb = COLOR_GREEN
    if title_en:
        p = doc.add_paragraph()
        run = p.add_run(title_en)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = COLOR_GRAY
        p.paragraph_format.space_after = Pt(12)


def add_sub_heading(doc: Document, title_ne: str, title_en: str = ""):
    heading = doc.add_heading(title_ne, level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 120, 0)
    if title_en:
        p = doc.add_paragraph()
        run = p.add_run(title_en)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = COLOR_GRAY


def add_body(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = COLOR_GRAY


def add_rationale_box(doc: Document, text_ne: str, text_en: str = ""):
    p = doc.add_paragraph()
    run = p.add_run("सिफारिसको आधार (Rationale): ")
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(139, 69, 19)
    if text_ne:
        run2 = p.add_run(text_ne)
        run2.font.size = Pt(11)
        run2.italic = True
    if text_en:
        p2 = doc.add_paragraph()
        run3 = p2.add_run(text_en)
        run3.font.size = Pt(10)
        run3.italic = True
        run3.font.color.rgb = COLOR_GRAY


def _add_image(doc: Document, buf: io.BytesIO, width_mm: float = IMG_WIDTH_MM):
    try:
        doc.add_picture(buf, width=Mm(width_mm))
    except Exception as e:
        logger.warning(f"Could not embed image: {e}")
        add_body(doc, "[Image unavailable]", italic=True, size=10)


def _add_table(doc: Document, headers: List[str], rows: List[List],
               col_widths: Optional[List[float]] = None):
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = COLOR_WHITE
        _set_cell_shading(cell, '006400')

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)


import logging
logger = logging.getLogger(__name__)


def build_chapter_1(doc: Document, basic_info: Dict, boundary: Dict):
    add_section_heading(doc, "१", "परिचय", "Introduction")

    narrative = get_chapter_narrative(1, basic_info=basic_info, boundary=boundary)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    blocks = boundary.get("blocks", [])
    if blocks:
        add_sub_heading(doc, "ब्लक विवरण", "Block Details")
        headers = ["क्र.सं.", "ब्लक नाम", "क्षेत्रफल (हेक्टर)"]
        rows = [
            [i + 1, b.get("name", ""), round(b.get("area_hectares", 0), 2)]
            for i, b in enumerate(blocks)
        ]
        _add_table(doc, headers, rows)

    add_body(doc, "यस योजनाको मुख्य उद्देश्य वनको दिगो व्यवस्थापन, "
                  "जैविक विविधता संरक्षण र उपभोक्ताको आवश्यकता पूर्ति गर्नु हो.")
    add_body(doc, "The main objective is sustainable forest management, "
                  "biodiversity conservation, and meeting community needs.", italic=True)


def build_chapter_2(doc: Document, db: Session, calc_id: UUID, basic_info: Dict,
                    boundary: Dict, raster: Dict, forest_name: str):
    add_section_heading(doc, "२", "भौगोलिक अवस्थिति", "Geographical Location")

    map_buf = generate_standard_map(db, calc_id, "boundary", forest_name=forest_name)
    _add_image(doc, map_buf)
    add_caption(doc, "Figure 2.1: सिमाना नक्सा — Boundary Map with Blocks")

    blocks_count = len(boundary.get("blocks", []))
    map_desc = describe_layer("boundary", raster, basic_info, forest_name, blocks_count)
    add_body(doc, map_desc["ne"])
    add_body(doc, map_desc["en"], italic=True)

    features = boundary.get("features", {})
    dir_labels = {"north": "उत्तर", "south": "दक्षिण", "east": "पूर्व", "west": "पश्चिम"}
    for direction, label_ne in dir_labels.items():
        items = features.get(direction, [])
        names = ", ".join([f.get("name", "") for f in items[:5]])
        if names:
            add_body(doc, f"{label_ne}: {names}", size=10)

    blocks = boundary.get("blocks", [])
    if blocks:
        add_sub_heading(doc, "ब्लक विवरण तालिका", "Block Details Table")
        headers = ["ब्लक", "क्षेत्रफल (हेक्टर)", "प्रभावकारी क्षेत्रफल (हेक्टर)"]
        rows = [
            [b.get("name", ""), round(b.get("area_hectares", 0), 2),
             round(b.get("effective_area_hectares", b.get("area_hectares", 0)), 2)]
            for b in blocks
        ]
        _add_table(doc, headers, rows)

    narrative = get_chapter_narrative(2, basic_info=basic_info, boundary=boundary)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_3(doc: Document, db: Session, calc_id: UUID,
                    raster: Dict, basic_info: Dict, forest_name: str):
    add_section_heading(doc, "३", "भौतिक वातावरण", "Physical Environment")

    maps_to_gen = [
        ("dem", "Figure 3.1: उचाइ नक्सा — Elevation Map"),
        ("slope", "Figure 3.2: भिरालो नक्सा — Slope Map"),
        ("aspect", "Figure 3.3: दिशा नक्सा — Aspect Map"),
        ("soil_texture", "Figure 3.4: माटो बनावट नक्सा — Soil Texture Map"),
    ]

    for layer, caption in maps_to_gen:
        try:
            map_buf = generate_standard_map(db, calc_id, layer, forest_name=forest_name)
            _add_image(doc, map_buf)
            add_caption(doc, caption)
        except Exception as e:
            logger.error(f"Map generation failed for {layer}: {e}")
            add_body(doc, f"[{caption} — map generation failed]", italic=True)

        map_desc = describe_layer(layer, raster, forest_name=forest_name)
        add_body(doc, map_desc["ne"])
        add_body(doc, map_desc["en"], italic=True)

    # Charts
    add_sub_heading(doc, "भौतिक वातावरण चार्टहरू", "Physical Environment Charts")

    slope_pcts = raster.get("slope_percentages", {})
    if slope_pcts:
        buf = generate_slope_pie(slope_pcts, forest_name=forest_name)
        if buf:
            _add_image(doc, buf)

    asp_pcts = raster.get("aspect_percentages", {})
    if asp_pcts:
        buf = generate_aspect_rose(asp_pcts, forest_name=forest_name)
        if buf:
            _add_image(doc, buf)

    narrative = get_chapter_narrative(3, raster=raster)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_4(doc: Document, db: Session, calc_id: UUID,
                    raster: Dict, species_data: Dict, forest_name: str):
    add_section_heading(doc, "४", "वन प्रकार तथा भू-आवरण", "Forest Type & Land Cover")

    maps_to_gen = [
        ("forest_type", "Figure 4.1: वन प्रकार नक्सा — Forest Type Map"),
        ("landcover", "Figure 4.2: भू-आवरण नक्सा — Land Cover Map"),
        ("canopy", "Figure 4.3: वन छाना नक्सा — Canopy Cover Map"),
    ]

    for layer, caption in maps_to_gen:
        try:
            map_buf = generate_standard_map(db, calc_id, layer, forest_name=forest_name)
            _add_image(doc, map_buf)
            add_caption(doc, caption)
        except Exception as e:
            logger.error(f"Map generation failed for {layer}: {e}")
            add_body(doc, f"[{caption} — map generation failed]", italic=True)

        map_desc = describe_layer(layer, raster, forest_name=forest_name)
        add_body(doc, map_desc["ne"])
        add_body(doc, map_desc["en"], italic=True)

    # Charts
    add_sub_heading(doc, "वन प्रकार तथा भू-आवरण चार्टहरू", "Forest Type & Land Cover Charts")

    ft_pcts = raster.get("forest_type_percentages", {})
    if ft_pcts:
        buf = generate_forest_type_pie(ft_pcts, forest_name=forest_name)
        if buf:
            _add_image(doc, buf)

    lc_pcts = raster.get("landcover_percentages", {})
    if lc_pcts:
        buf = generate_landcover_pie(lc_pcts, forest_name=forest_name)
        if buf:
            _add_image(doc, buf)

    narrative = get_chapter_narrative(4, raster=raster, species_data=species_data)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_5(doc: Document, db: Session, calc_id: UUID,
                    raster: Dict, biodiversity: Dict, forest_name: str):
    add_section_heading(doc, "५", "वन स्वास्थ्य तथा जैविक विविधता",
                        "Forest Health & Biodiversity")

    maps_to_gen = [
        ("forest_health", "Figure 5.1: वन स्वास्थ्य नक्सा — Forest Health Map"),
        ("biomass", "Figure 5.2: बायोमास नक्सा — Biomass Map"),
    ]

    for layer, caption in maps_to_gen:
        try:
            map_buf = generate_standard_map(db, calc_id, layer, forest_name=forest_name)
            _add_image(doc, map_buf)
            add_caption(doc, caption)
        except Exception as e:
            logger.error(f"Map generation failed for {layer}: {e}")
            add_body(doc, f"[{caption} — map generation failed]", italic=True)

        map_desc = describe_layer(layer, raster, forest_name=forest_name)
        add_body(doc, map_desc["ne"])
        add_body(doc, map_desc["en"], italic=True)

    fh_pcts = raster.get("forest_health_percentages", {})
    if fh_pcts:
        add_sub_heading(doc, "वन स्वास्थ्य वितरण", "Forest Health Distribution")
        headers = ["स्वास्थ्य अवस्था", "प्रतिशत"]
        rows = [[k, f"{v:.1f}%"] for k, v in sorted(fh_pcts.items(), key=lambda x: x[1], reverse=True)]
        _add_table(doc, headers, rows)

    narrative = get_chapter_narrative(5, raster=raster, biodiversity=biodiversity)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_6(doc: Document, mgmt_data: Dict, forest_name: str):
    add_section_heading(doc, "६", "वन स्रोत सर्वेक्षण", "Forest Resource Survey")

    # Species composition
    sp_data = mgmt_data.get("species_composition", {})
    fw = sp_data.get("forest_wide", [])
    if fw:
        add_sub_heading(doc, "प्रजाति संरचना", "Species Composition")
        buf = chart_species_composition(sp_data, forest_name=forest_name)
        _add_image(doc, buf)
        add_caption(doc, "Figure 6.1: प्रजाति संरचना — Species Composition")

        headers = ["प्रजाति", "स्थानीय नाम", "आयतन (m³/ha)", "प्रतिशत", "गणना (N/ha)"]
        rows = [
            [s.get("scientific_name", "")[:25], s.get("local_name", ""),
             round(s.get("total_volume_m3_per_ha", 0), 2),
             round(s.get("volume_pct", 0), 1),
             round(s.get("total_count_per_ha", 0), 1)]
            for s in fw[:10]
        ]
        _add_table(doc, headers, rows)

    # Block comparison
    bc_data = mgmt_data.get("block_comparison", {})
    ranked = bc_data.get("ranked", [])
    if ranked:
        add_sub_heading(doc, "ब्लक तुलना", "Block Comparison")
        buf = chart_block_comparison(bc_data, forest_name=forest_name)
        _add_image(doc, buf)
        add_caption(doc, "Figure 6.2: ब्लक तुलना — Block Comparison by Growing Stock")

        headers = ["क्र.म.", "ब्लक", "क्षेत्रफल (ha)", "Growing Stock (m³/ha)"]
        rows = [
            [r.get("rank", ""), r.get("name", ""), round(r.get("area_ha", 0), 2),
             round(r.get("growing_stock_m3ha", 0), 2)]
            for r in ranked
        ]
        _add_table(doc, headers, rows)

    # Carbon
    carb = mgmt_data.get("carbon_per_block", {})
    carb_blocks = carb.get("blocks", [])
    if carb_blocks:
        add_sub_heading(doc, "कार्बन भण्डार", "Carbon Stock")
        buf = chart_carbon_stock(carb, forest_name=forest_name)
        _add_image(doc, buf)
        add_caption(doc, "Figure 6.3: कार्बन भण्डार — Carbon Stock per Block")

    # DBH class volume
    dbh = mgmt_data.get("dbh_class_volume", {})
    dbh_blocks = dbh.get("blocks", [])
    if dbh_blocks:
        add_sub_heading(doc, "DBH वर्ग आयतन", "DBH Class Volume Distribution")
        buf = chart_dbh_class_volume(dbh, forest_name=forest_name)
        _add_image(doc, buf)
        add_caption(doc, "Figure 6.4: DBH वर्ग आयतन — DBH Class Volume")

    narrative = get_chapter_narrative(6, mgmt_data=mgmt_data)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_7(doc: Document, mgmt_data: Dict, forest_name: str):
    add_section_heading(doc, "७", "वार्षिक स्वीकार्य कटान (AAH)",
                        "Annual Allowable Harvest")

    bc = mgmt_data.get("block_comparison", {})
    ranked = bc.get("ranked", [])
    if ranked:
        add_sub_heading(doc, "AAH ब्लक तुलना", "AAH Block Comparison")
        buf = chart_block_comparison(bc, forest_name=forest_name)
        _add_image(doc, buf)
        add_caption(doc, "Figure 7.1: AAH ब्लक तुलना — AAH Block Comparison")

        headers = ["ब्लक", "अवस्था", "क्षेत्रफल (ha)",
                   "AAH काठ (m³/yr)", "AAH दाउरा (m³/yr)"]
        rows = [
            [r.get("name", ""), r.get("condition", ""),
             round(r.get("area_ha", 0), 2),
             round(r.get("aah_timber_m3yr", 0), 2),
             round(r.get("aah_fuelwood_m3yr", 0), 2)]
            for r in ranked
        ]
        _add_table(doc, headers, rows)

    # Growth rate
    gr = mgmt_data.get("growth_rate_classification", {})
    if gr.get("classes"):
        add_sub_heading(doc, "वृद्धि दर वर्गीकरण", "Growth Rate Classification")
        buf = chart_growth_rate(gr, forest_name=forest_name)
        _add_image(doc, buf)
        add_caption(doc, "Figure 7.2: वृद्धि दर — Growth Rate Classification")

    # Productivity
    prod = mgmt_data.get("productivity_classification", {})
    if prod.get("classes"):
        add_sub_heading(doc, "उत्पादकता वर्गीकरण", "Productivity Classification")
        buf = chart_productivity(prod, forest_name=forest_name)
        _add_image(doc, buf)
        add_caption(doc, "Figure 7.3: उत्पादकता — Productivity Classification")

    narrative = get_chapter_narrative(7, mgmt_data=mgmt_data)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_8(doc: Document, db: Session, calc_id: UUID,
                    mgmt_data: Dict, ten_year: Dict, forest_name: str):
    add_section_heading(doc, "८", "१० वर्षे ब्लक योजना", "10-Year Block Plan")

    try:
        map_buf = generate_standard_map(db, calc_id, "boundary", forest_name=forest_name)
        _add_image(doc, map_buf)
        add_caption(doc, "Figure 8.1: ब्लक नक्सा — Block Map for 10-Year Plan")
    except Exception as e:
        logger.error(f"Block map generation failed: {e}")
        add_body(doc, "[Block map unavailable]", italic=True)

    block_schedule = ten_year.get("block_schedule", {})
    if block_schedule:
        headers = ["ब्लक", "अवस्था", "क्षेत्रफल (ha)",
                   "Growing Stock (m³/ha)", "AAH (m³/yr)",
                   "फसल वर्ष", "घुमाउ (yrs)"]
        rows = []
        for name, sched in sorted(block_schedule.items()):
            rows.append([
                name, sched.get("condition", ""),
                round(sched.get("area_ha", 0), 2),
                round(sched.get("growing_stock_m3ha", 0), 2),
                round(sched.get("aah_timber_m3yr", 0), 2),
                ", ".join(str(y) for y in sched.get("harvest_years", [])),
                sched.get("rotation_yrs", 0),
            ])

        add_sub_heading(doc, "ब्लक योजना सारांश", "Block Plan Summary")
        _add_table(doc, headers, rows)

        # Per-block rationale
        for name, sched in sorted(block_schedule.items()):
            rationale = sched.get("rationale", "")
            rationale_en = sched.get("rationale_en", "")
            add_rationale_box(doc,
                f"ब्लक '{name}' — {rationale}",
                f"Block '{name}' — {rationale_en}")

    summary = ten_year.get("summary", {})
    if summary:
        add_sub_heading(doc, "योजना अवधि सारांश", "Plan Period Summary")
        headers = ["विवरण", "मान"]
        rows = [
            ["जम्मा कटान (m³)", str(round(summary.get("total_harvest_m3_10yr", 0), 2))],
            ["जम्मा बजेट (रु.)", f"{summary.get('total_budget_10yr', 0):,.0f}"],
            ["वार्षिक औसत कटान (m³)", str(round(summary.get("average_yearly_harvest_m3", 0), 2))],
            ["ब्लक संख्या", str(summary.get("total_blocks", 0))],
        ]
        _add_table(doc, headers, rows)

    narrative = get_chapter_narrative(8, mgmt_data=mgmt_data, ten_year=ten_year)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_9(doc: Document, mgmt_data: Dict, raster: Dict, forest_name: str):
    add_section_heading(doc, "९", "वन संवर्द्धन तथा संरक्षण",
                        "Forest Promotion & Conservation")

    fc = mgmt_data.get("forest_condition_summary", {})
    if fc.get("by_condition") or fc.get("regeneration"):
        add_sub_heading(doc, "वन अवस्था तथा पुनरुत्पादन",
                        "Forest Condition & Regeneration")
        buf = chart_forest_condition(fc, forest_name=forest_name)
        _add_image(doc, buf)
        add_caption(doc, "Figure 9.1: वन अवस्था — Forest Condition & Regeneration")

    regen = fc.get("regeneration", [])
    if regen:
        headers = ["ब्लक", "अवस्था", "बिरुवा (N/ha)", "लाथ्रा (N/ha)", "जम्मा (N/ha)"]
        rows = [
            [r.get("block", ""), r.get("condition", ""),
             round(r.get("seedling_nha", 0), 1),
             round(r.get("sapling_nha", 0), 1),
             round(r.get("total_nha", 0), 1)]
            for r in regen
        ]
        _add_table(doc, headers, rows)

    narrative = get_chapter_narrative(9, raster=raster, mgmt_data=mgmt_data)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_10(doc: Document, activities: Dict, ten_year: Dict, mgmt_data: Dict,
                     db: Session, calc_id: UUID, forest_name: str):
    add_section_heading(doc, "१०", "वार्षिक क्रियाकलाप तथा बजेट",
                        "Annual Activities & Budget")

    summary = ten_year.get("summary", {})
    block_schedule = ten_year.get("block_schedule", {})

    # Year-by-year activity table
    years_data = ten_year.get("years", {})
    if years_data:
        add_sub_heading(doc, "वर्ष अनुसार क्रियाकलाप", "Year-wise Activities")
        for y in range(1, 11):
            acts = years_data.get(y, [])
            if acts:
                add_body(doc, f"वर्ष {y}:", bold=True, size=11)
                headers = ["ब्लक", "क्रियाकलाप", "क्षेत्रफल (ha)",
                           "काठ (m³)", "बजेट (रु.)"]
                rows = [
                    [a.get("block", ""), a.get("activity", ""),
                     round(a.get("area_ha", 0), 2),
                     round(a.get("harvest_timber_m3", 0), 2),
                     f"{a.get('budget', 0):,.0f}"]
                    for a in acts
                ]
                _add_table(doc, headers, rows)

    narrative = get_chapter_narrative(10, activities=activities, ten_year=ten_year)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_11(doc: Document, activities: Dict, ten_year: Dict):
    add_section_heading(doc, "११", "वित्तीय विश्लेषण", "Financial Analysis")

    summary = ten_year.get("summary", {})
    total_budget = summary.get("total_budget_10yr", 0)
    avg_budget = summary.get("average_yearly_budget", 0)

    headers = ["विवरण", "रकम (रु.)"]
    rows = [
        ["योजना अवधिको जम्मा बजेट", f"{total_budget:,.0f}"],
        ["वार्षिक औसत बजेट", f"{avg_budget:,.0f}"],
    ]

    if activities.get("available"):
        budget = activities.get("total_budget", 0)
        rows.append(["क्रियाकलाप बजेट", f"{budget:,.0f}"])

    _add_table(doc, headers, rows)

    narrative = get_chapter_narrative(11, activities=activities)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    doc.add_page_break()


def build_chapter_12(doc: Document):
    add_section_heading(doc, "१२", "अनुगमन तथा मूल्याङ्कन",
                        "Monitoring & Evaluation")

    narrative = get_chapter_narrative(12)
    add_body(doc, narrative["ne"])
    add_body(doc, narrative["en"], italic=True)

    headers = ["सूचक", "विधि", "आवृत्ति", "जिम्मेवार"]
    rows = [
        ["कटान मात्रा", "क्षेत्रगत निरीक्षण", "६ महिना", "वन समिति"],
        ["पुनरुत्पादन", "नमुना प्लट", "वार्षिक", "प्राविधिक सहायक"],
        ["बजेट खर्च", "लेखा परीक्षण", "त्रैमासिक", "कोषाध्यक्ष"],
        ["वन स्वास्थ्य", "अवलोकन", "वार्षिक", "वन समिति"],
        ["उपभोक्ता सन्तुष्टि", "सर्वेक्षण", "२ वर्ष", "वन समिति"],
    ]
    _add_table(doc, headers, rows)

    add_body(doc, "योजना प्रभावकारी कार्यान्वयनको लागि नियमित अनुगमन "
                  "अत्यावश्यक छ। प्रत्येक ६ महिनामा समीक्षा बैठक गरी "
                  "आवश्यक सुधार गर्नुपर्दछ।")
    add_body(doc, "Regular monitoring is essential for effective implementation. "
                  "Review meetings should be held every 6 months.", italic=True)
