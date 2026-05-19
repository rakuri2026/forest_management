import io
import logging
from typing import Dict, Any, Optional
from uuid import UUID
from datetime import datetime
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from . import plan_chapter_builders as chapters
from .plan_ten_year import build_ten_year_plan
from ..report.data_collector import (
    collect_all_data,
    get_calculation_basic_info,
    get_boundary_info,
    get_species_info,
    get_raster_analysis,
    get_biodiversity_data,
    get_activities_data,
    get_household_data,
    get_committee_data,
)
from ..field_inventory_mgmt_data import get_management_plan_data
from ..field_inventory_mgmt_charts import chart_species_composition, chart_block_comparison

logger = logging.getLogger(__name__)

COLOR_GREEN = RGBColor(0, 100, 0)
COLOR_DARK = RGBColor(33, 33, 33)
COLOR_GRAY = RGBColor(128, 128, 128)


def _build_cover_page(doc: Document, forest_name: str, basic_info: Dict,
                      ten_year_summary: Dict):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{forest_name}")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = COLOR_GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("सामुदायिक वन व्यवस्थापन योजना")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = COLOR_GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Community Forest 10-Year Management Plan")
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()

    district = basic_info.get("district", "———")
    municipality = basic_info.get("municipality", "———")
    ward = basic_info.get("ward", "———")
    prov = basic_info.get("province", "———")
    area = basic_info.get("total_area_hectares", 0)

    details = [
        ("प्रदेश / जिल्ला / पालिका", f"प्रदेश {prov} / {district} / {municipality}–{ward}"),
        ("क्षेत्रफल", f"{area:.2f} हेक्टर"),
        ("योजना अवधि", "आ.व. २०८२/८३ देखि २०९१/९२ सम्म (१० वर्ष)"),
        ("Plan Period", "FY 2082/83 to FY 2091/92 (10 Years)"),
    ]

    table = doc.add_table(rows=len(details), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(details):
        cell_l = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_l.text = label
        cell_v.text = str(value)
        for cell in [cell_l, cell_v]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    if cell == cell_l:
                        run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"मिति: {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_GRAY

    doc.add_page_break()


def _add_toc(doc: Document):
    heading = doc.add_heading('विषय सूची (Table of Contents)', level=1)
    for run in heading.runs:
        run.font.color.rgb = COLOR_GREEN

    p = doc.add_paragraph()
    run = p.add_run(
        "[तलको ठाउँमा विषय सूची स्वचालित रूपमा उत्पन्न हुनेछ। "
        "MS Word मा Ctrl+A थिची F9 थिच्नुहोस्।]"
    )
    run.font.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(10)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "[Table of Contents will be auto-generated. "
        "In MS Word, press Ctrl+A then F9 to update.]"
    )
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(150, 150, 150)
    run2.font.size = Pt(10)

    doc.add_page_break()


def _add_annexes(doc: Document, species_data: Dict, household_data: Dict,
                 committee_data: Dict, biodiversity: Dict):
    chapters.add_section_heading(doc, "परिशिष्ट", "Annexes")

    # Annex 1: Species list
    chapters.add_sub_heading(doc, "परिशिष्ट १: प्रजाति सूची", "Annex 1: Species List")
    species_list = species_data.get("species_list", [])
    if species_list:
        headers = ["वैज्ञानिक नाम", "स्थानीय नाम", "वृद्धि दर", "आर्थिक मूल्य"]
        rows = [
            [s.get("scientific_name", "")[:30], s.get("local_name", ""),
             s.get("growth_rate", ""), s.get("economic_value", "")]
            for s in species_list[:20]
        ]
        chapters._add_table(doc, headers, rows)

    # Annex 2: Household info
    chapters.add_sub_heading(doc, "परिशिष्ट २: घरधुरी विवरण", "Annex 2: Household Details")
    if household_data.get("available"):
        hh = household_data
        headers = ["विवरण", "मान"]
        rows = [
            ["जम्मा घरधुरी", str(hh.get("total_households", 0))],
            ["जम्मा जनसंख्या", str(hh.get("total_population", 0))],
            ["काठ माग (cft)", str(round(hh.get("timber_demand_cft", 0), 2))],
            ["दाउरा माग (भारी)", str(round(hh.get("firewood_demand_bhari", 0), 2))],
        ]
        chapters._add_table(doc, headers, rows)

    # Annex 3: Committee
    chapters.add_sub_heading(doc, "परिशिष्ट ३: समिति विवरण", "Annex 3: Committee Details")
    uc = committee_data.get("user_committee", {})
    if uc.get("members"):
        headers = ["नाम", "पद", "लिङ्ग", "ठेगाना"]
        rows = [
            [m.get("name", ""), m.get("position", ""),
             m.get("gender", ""), m.get("address", "")]
            for m in uc.get("members", [])
        ]
        chapters._add_table(doc, headers, rows)

    doc.add_page_break()


def generate_10yr_management_plan_docx(
    db: Session,
    field_inventory_id: UUID,
    calculation_id: UUID,
    aah_good: float = 75.0,
    aah_moderate: float = 60.0,
    aah_weak: float = 40.0,
    include_maps: bool = True,
    include_charts: bool = True,
) -> bytes:
    """
    Generate a complete 12-chapter, 10-year management plan DOCX.

    Args:
        db: Database session
        field_inventory_id: Field inventory UUID
        calculation_id: Forest calculation UUID
        aah_good: AAH multiplier for Good condition (%)
        aah_moderate: AAH multiplier for Moderate condition (%)
        aah_weak: AAH multiplier for Weak condition (%)
        include_maps: Whether to include map images
        include_charts: Whether to include chart images

    Returns:
        DOCX file as bytes
    """
    from ...models.calculation import Calculation
    from ...models.field_inventory import FieldInventoryCalculation

    fi = db.query(FieldInventoryCalculation).filter(
        FieldInventoryCalculation.id == field_inventory_id
    ).first()
    if not fi:
        raise ValueError("Field inventory not found")

    calc = db.query(Calculation).filter(Calculation.id == calculation_id).first()
    forest_name = calc.forest_name if calc else "Unknown Forest"

    # ── Collect ALL data ──
    logger.info("Collecting management plan data...")
    mgmt_data = get_management_plan_data(
        db, field_inventory_id, calculation_id,
        aah_good, aah_moderate, aah_weak,
    )

    logger.info("Collecting report data...")
    basic_info = get_calculation_basic_info(db, str(calculation_id))
    boundary = get_boundary_info(db, str(calculation_id))
    species_data = get_species_info(db, str(calculation_id))
    raster_data = get_raster_analysis(db, str(calculation_id))
    biodiversity = get_biodiversity_data(db, str(calculation_id))
    activities = get_activities_data(db, str(calculation_id))
    household_data = get_household_data(db, str(calculation_id))
    committee_data = get_committee_data(db, str(calculation_id))

    # ── Build 10-year plan ──
    logger.info("Building 10-year plan...")
    ten_year = build_ten_year_plan(mgmt_data, activities)

    # ── Create Document ──
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cover page
    _build_cover_page(doc, forest_name, basic_info, ten_year.get("summary", {}))

    # TOC
    _add_toc(doc)

    # ── Build Chapters ──
    logger.info("Building chapter 1: Introduction")
    chapters.build_chapter_1(doc, basic_info, boundary)

    if include_maps:
        logger.info("Building chapter 2: Geographical Location")
        chapters.build_chapter_2(doc, db, calculation_id, basic_info, boundary,
                                  raster_data, forest_name)

        logger.info("Building chapter 3: Physical Environment")
        chapters.build_chapter_3(doc, db, calculation_id, raster_data,
                                  basic_info, forest_name)

        logger.info("Building chapter 4: Forest Type & Land Cover")
        chapters.build_chapter_4(doc, db, calculation_id, raster_data,
                                  species_data, forest_name)

        logger.info("Building chapter 5: Forest Health & Biodiversity")
        chapters.build_chapter_5(doc, db, calculation_id, raster_data,
                                  biodiversity, forest_name)

    if include_charts:
        logger.info("Building chapter 6: Forest Resource Survey")
        chapters.build_chapter_6(doc, mgmt_data, forest_name)

        logger.info("Building chapter 7: AAH")
        chapters.build_chapter_7(doc, mgmt_data, forest_name)

    logger.info("Building chapter 8: 10-Year Block Plan")
    chapters.build_chapter_8(doc, db, calculation_id, mgmt_data, ten_year, forest_name)

    logger.info("Building chapter 9: Forest Promotion & Conservation")
    chapters.build_chapter_9(doc, mgmt_data, raster_data, forest_name)

    logger.info("Building chapter 10: Annual Activities & Budget")
    chapters.build_chapter_10(doc, activities, ten_year, mgmt_data,
                               db, calculation_id, forest_name)

    logger.info("Building chapter 11: Financial Analysis")
    chapters.build_chapter_11(doc, activities, ten_year)

    logger.info("Building chapter 12: Monitoring & Evaluation")
    chapters.build_chapter_12(doc)

    # Annexes
    logger.info("Building annexes...")
    _add_annexes(doc, species_data, household_data, committee_data, biodiversity)

    # ── Save ──
    logger.info("Saving document...")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
