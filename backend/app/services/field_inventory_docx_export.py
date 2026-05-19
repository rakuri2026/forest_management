"""
Management Plan DOCX Export
Builds a comprehensive DOCX document with embedded maps, charts, and data tables.
"""
import io
import logging
from typing import Dict, Any, Optional
from uuid import UUID
from datetime import datetime
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Inches, Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from .field_inventory_mgmt_data import get_management_plan_data

# Ensure all models are registered for SQLAlchemy mapper configuration
from ..models import yearly_activities  # noqa: F401 — registers ProposedYearlyActivity model
from .field_inventory_mgmt_maps import generate_map_image
from .field_inventory_mgmt_charts import (
    chart_species_composition,
    chart_block_comparison,
    chart_annual_harvest,
    chart_forest_condition,
    chart_dbh_class_volume,
    chart_carbon_stock,
    chart_growth_rate,
    chart_stand_structure,
    chart_productivity,
)

logger = logging.getLogger(__name__)

DOCX_IMG_WIDTH_MM = 150
DOCX_PAGE_WIDTH_MM = 210
MAP_GEN_DPI = 200
COLOR_GREEN = RGBColor(46, 125, 50)
COLOR_DARK = RGBColor(33, 33, 33)
COLOR_GRAY = RGBColor(128, 128, 128)


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def _add_table(doc, headers: list, rows: list, col_widths: Optional[list] = None):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, '2F5496')

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    return table


def _add_section_heading(doc, text: str, level: int = 1):
    """Add a Nepali section heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = COLOR_GREEN


def _add_body_text(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    """Add a paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def _add_image_safe(doc, buf: io.BytesIO, width_mm: float = DOCX_IMG_WIDTH_MM):
    """Add image at given width in mm, with fallback on error."""
    try:
        doc.add_picture(buf, width=Mm(width_mm))
    except Exception as e:
        logger.warning(f"Could not embed image: {e}")
        _add_body_text(doc, "[Image could not be embedded]", italic=True, size=10)


def _get_forest_info(db: Session, calculation_id: UUID) -> Dict[str, Any]:
    """Get forest name, district, etc."""
    from ..models.calculation import Calculation
    calc = db.query(Calculation).filter(Calculation.id == calculation_id).first()
    if not calc:
        return {"name": "Unknown Forest", "area_ha": 0}
    rd = calc.result_data or {}
    return {
        "name": calc.forest_name or "Unknown Forest",
        "area_ha": float(rd.get("area_hectares", 0)),
        "district": rd.get("whole_district", ""),
        "municipality": rd.get("whole_municipality", ""),
        "ward": rd.get("whole_ward", ""),
    }


MAP_LAYERS = [
    ("boundary", "सिमाना नक्सा", "Boundary Map"),
    ("forest_type", "वन प्रकार नक्सा", "Forest Type Map"),
    ("forest_health", "वन स्वास्थ्य नक्सा", "Forest Health Map"),
    ("slope", "भिरालो नक्सा", "Slope Map"),
]

CHART_SECTIONS = [
    ("species_composition", chart_species_composition, "प्रजाती संरचना", "Species Composition"),
    ("block_comparison", chart_block_comparison, "ब्लक तुलना", "Block Comparison"),
    ("annual_harvest_plan", chart_annual_harvest, "वार्षिक फसल योजना", "Annual Harvest Plan"),
    ("forest_condition_summary", chart_forest_condition, "वन स्थिति सारांश", "Forest Condition"),
    ("dbh_class_volume", chart_dbh_class_volume, "DBH वर्ग आयतन", "DBH Class Volume"),
    ("carbon_per_block", chart_carbon_stock, "कार्बन भण्डार", "Carbon Stock"),
    ("growth_rate_classification", chart_growth_rate, "वृद्धि दर वर्गीकरण", "Growth Rate"),
    ("stand_structure", chart_stand_structure, "रुख संरचना प्रोफाइल", "Stand Structure"),
    ("productivity_classification", chart_productivity, "उत्पादकता वर्गीकरण", "Productivity"),
]


def generate_management_plan_docx(
    db: Session,
    field_inventory_id: UUID,
    calculation_id: UUID,
    aah_good: float = 75.0,
    aah_moderate: float = 60.0,
    aah_weak: float = 40.0,
) -> bytes:
    from ..models.field_inventory import FieldInventoryCalculation
    fi = db.query(FieldInventoryCalculation).filter(
        FieldInventoryCalculation.id == field_inventory_id
    ).first()
    if not fi:
        raise ValueError("Field inventory not found")

    forest_info = _get_forest_info(db, calculation_id)
    forest_name = forest_info["name"]

    mgmt_data = get_management_plan_data(db, field_inventory_id, calculation_id, aah_good, aah_moderate, aah_weak)

    doc = Document()

    # ── COVER PAGE ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n")
    run.font.size = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{forest_name}\nव्यवस्थापन योजना प्रतिवेदन")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = COLOR_GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Management Plan Report")
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n{forest_info.get('district', '')} जिल्ला, {forest_info.get('municipality', '')}")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    area = forest_info.get("area_ha", 0)
    run = p.add_run(f"क्षेत्रफल: {area:.2f} हेक्टर")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\nमिति: {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_GRAY

    doc.add_page_break()

    # ── SECTION 1: MAPS ──
    _add_section_heading(doc, "भाग १: नक्साहरू", level=1)
    _add_body_text(doc, "Section 1: Maps", italic=True, size=10)

    for layer_key, ne_label, en_label in MAP_LAYERS:
        _add_section_heading(doc, f"{ne_label} ({en_label})", level=2)
        try:
            map_buf = generate_map_image(db, calculation_id, layer_key, dpi=MAP_GEN_DPI, forest_name=forest_name)
            _add_image_safe(doc, map_buf)
        except Exception as e:
            logger.error(f"Failed to generate {layer_key} map: {e}")
            _add_body_text(doc, f"[{ne_label} — map generation failed]", italic=True)
        doc.add_page_break()

    # ── SECTION 2: MANAGEMENT CHARTS ──
    _add_section_heading(doc, "भाग २: व्यवस्थापन चार्टहरू", level=1)
    _add_body_text(doc, "Section 2: Management Charts & Tables", italic=True, size=10)

    for data_key, chart_fn, ne_label, en_label in CHART_SECTIONS:
        _add_section_heading(doc, f"{ne_label} ({en_label})", level=2)
        chart_data = mgmt_data.get(data_key, {})

        # Chart image
        try:
            chart_buf = chart_fn(chart_data, forest_name=forest_name)
            _add_image_safe(doc, chart_buf)
        except Exception as e:
            logger.error(f"Failed to generate {data_key} chart: {e}")
            _add_body_text(doc, f"[{ne_label} — chart generation failed]", italic=True)

        # Data table (varies by section)
        try:
            _add_data_table_for_section(doc, data_key, chart_data)
        except Exception as e:
            logger.error(f"Failed to add {data_key} table: {e}")

        doc.add_page_break()

    # ── SAVE ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_data_table_for_section(doc: Document, data_key: str, data: Dict):
    """Add data table below each chart based on section type."""
    if data_key == "species_composition":
        fw = data.get("forest_wide", [])
        if fw:
            headers = ["Species", "Local Name", "Vol (m³/ha)", "Vol %", "Count (N/ha)"]
            rows = [
                [s.get("scientific_name", "")[:25], s.get("local_name", ""),
                 s.get("total_volume_m3_per_ha", 0), s.get("volume_pct", 0),
                 s.get("total_count_per_ha", 0)]
                for s in fw
            ]
            _add_table(doc, headers, rows)

    elif data_key == "block_comparison":
        ranked = data.get("ranked", [])
        if ranked:
            headers = ["Rank", "Block", "Area (ha)", "Growing Stock (m³/ha)", "AAH Timber (m³/yr)", "Condition"]
            rows = [
                [r.get("rank", ""), r.get("name", ""), r.get("area_ha", 0),
                 r.get("growing_stock_m3ha", 0), r.get("aah_timber_m3yr", 0),
                 r.get("condition", "")]
                for r in ranked
            ]
            _add_table(doc, headers, rows)

    elif data_key == "annual_harvest_plan":
        blocks = data.get("blocks", [])
        if blocks:
            headers = ["Block", "Area (ha)", "AAH Timber (m³/yr)", "AAH Fuelwood (m³/yr)",
                       "Coupe Area (ha)", "Rotation (yrs)", "Condition"]
            rows = [
                [b.get("name", ""), b.get("area_ha", 0), b.get("aah_timber_m3yr", 0),
                 b.get("aah_fuelwood_m3yr", 0), b.get("coupe_area_ha", 0),
                 b.get("rotation_yrs", 0), b.get("condition", "")]
                for b in blocks
            ]
            _add_table(doc, headers, rows)

    elif data_key == "forest_condition_summary":
        regen = data.get("regeneration", [])
        if regen:
            headers = ["Block", "Condition", "Seedling (N/ha)", "Sapling (N/ha)", "Total (N/ha)"]
            rows = [
                [r.get("block", ""), r.get("condition", ""), r.get("seedling_nha", 0),
                 r.get("sapling_nha", 0), r.get("total_nha", 0)]
                for r in regen
            ]
            _add_table(doc, headers, rows)

    elif data_key == "dbh_class_volume":
        blocks = data.get("blocks", [])
        if blocks:
            headers = ["Block", "DBH Class", "Count (N/ha)", "Timber (m³/ha)", "Fuelwood (m³/ha)", "Total (m³/ha)"]
            rows = []
            for b in blocks[:3]:
                blk = b.get("block", "")
                for cls in b.get("classes", []):
                    rows.append([
                        blk, cls.get("dbh_class", ""), cls.get("count_nha", 0),
                        cls.get("timber_m3ha", 0), cls.get("fuelwood_m3ha", 0),
                        cls.get("total_m3ha", 0),
                    ])
            if rows:
                _add_table(doc, headers, rows)

    elif data_key == "carbon_per_block":
        blocks = data.get("blocks", [])
        if blocks:
            headers = ["Block", "Area (ha)", "AGB (t/ha)", "BGB (t/ha)", "C Stock (t C/ha)", "CO₂e (t/ha)"]
            rows = [
                [b.get("block", ""), b.get("area_ha", 0), b.get("agb_tha", 0),
                 b.get("bgb_tha", 0), b.get("c_stock_tcha", 0), b.get("co2e_tha", 0)]
                for b in blocks
            ]
            _add_table(doc, headers, rows)

    elif data_key == "growth_rate_classification":
        classes = data.get("classes", [])
        if classes:
            headers = ["Growth Rate", "Species Count", "Volume (m³/ha)", "Volume %", "Key Species"]
            rows = [
                [c.get("rate", ""), c.get("species_count", 0), c.get("volume_m3_per_ha", 0),
                 c.get("volume_pct", 0), c.get("species", "")]
                for c in classes
            ]
            _add_table(doc, headers, rows)

    elif data_key == "stand_structure":
        blocks = data.get("blocks", [])
        if blocks:
            headers = ["Block", "DBH Class", "Actual (N/ha)", "Ideal (N/ha)", "Diff", "Status"]
            rows = []
            for b in blocks[:2]:
                blk = b.get("block", "")
                for cls in b.get("classes", []):
                    rows.append([
                        blk, cls.get("dbh_class", ""), cls.get("actual_nha", 0),
                        cls.get("ideal_nha", 0), cls.get("difference", 0),
                        cls.get("status", ""),
                    ])
            if rows:
                _add_table(doc, headers, rows)

    elif data_key == "productivity_classification":
        classes = data.get("classes", [])
        if classes:
            headers = ["Productivity", "Threshold", "Block Count", "Area (ha)", "Volume (m³)"]
            rows = [
                [c.get("class", ""), c.get("threshold", ""), c.get("block_count", 0),
                 c.get("area_ha", 0), c.get("volume_m3", 0)]
                for c in classes
            ]
            _add_table(doc, headers, rows)
