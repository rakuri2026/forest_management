"""
SVG utilities for embedding in DOCX
Embeds SVG directly via OPC archive manipulation (bypasses Cairo dep)
"""
import io
from typing import Optional
from base64 import b64decode
from lxml import etree


def svg_to_png_bytes(svg_bytes: bytes, scale: float = 2.0) -> bytes:
    """Convert SVG bytes to PNG bytes using CairoSVG"""
    import cairosvg
    return cairosvg.svg2png(bytestring=svg_bytes, scale=scale)


def svg_to_png_stream(svg_bytes: bytes, scale: float = 2.0) -> io.BytesIO:
    """Convert SVG bytes to PNG BytesIO"""
    png_bytes = svg_to_png_bytes(svg_bytes, scale)
    return io.BytesIO(png_bytes)


def is_svg_data_uri(data_uri: str) -> bool:
    """Check if a data URI is SVG"""
    return data_uri.startswith("data:image/svg+xml")


def ensure_png_stream(data_uri: str, scale: float = 2.0) -> Optional[io.BytesIO]:
    """
    Convert a data URI (SVG or PNG) to a PNG BytesIO suitable for python-docx.
    Falls back to None if CairoSVG is unavailable — callers should handle this.
    """
    if not data_uri or not data_uri.startswith("data:"):
        return None
    encoded = data_uri.split(",", 1)[1]
    raw_bytes = b64decode(encoded)
    if data_uri.startswith("data:image/svg+xml"):
        try:
            return svg_to_png_stream(raw_bytes, scale)
        except Exception:
            return None
    return io.BytesIO(raw_bytes)


def _decode_data_uri(data_uri: str) -> tuple[bytes, str]:
    """Decode a data URI into (raw_bytes, mime_type)."""
    if not data_uri or not data_uri.startswith("data:"):
        raise ValueError("Not a data URI")
    header, encoded = data_uri.split(",", 1)
    mime = header.split(";")[0].replace("data:", "")
    raw_bytes = b64decode(encoded)
    return raw_bytes, mime


def _get_svg_dimensions(svg_bytes: bytes) -> tuple[float, float]:
    """Parse SVG to get intrinsic width/height in pixels (or viewBox fallback)."""
    try:
        root = etree.fromstring(svg_bytes)
        vb = root.get("viewBox")
        if vb:
            parts = vb.strip().split()
            if len(parts) == 4:
                return float(parts[2]), float(parts[3])
        for attr in ("width", "height"):
            val = root.get(attr)
            if val:
                import re
                numeric = re.sub(r"[^\d.]", "", val)
                if numeric:
                    w = float(re.sub(r"[^\d.]", "", root.get("width", "600")))
                    h = float(re.sub(r"[^\d.]", "", root.get("height", "400")))
                    return w, h
    except Exception:
        pass
    return 600.0, 400.0


def add_svg_picture(doc, data_uri_or_bytes, width_inches=5.0):
    """
    Embed an SVG image directly into the DOCX via OPC archive manipulation.
    Bypasses python-docx's add_picture() which rejects SVG.

    Args:
        doc: python-docx Document
        data_uri_or_bytes: "data:image/svg+xml;base64,..." string or raw SVG bytes
        width_inches: desired display width in inches (default 5.0)
    Returns:
        The added picture shape, or None if embedding fails.
    """
    if isinstance(data_uri_or_bytes, str) and data_uri_or_bytes.startswith("data:"):
        raw_bytes, _ = _decode_data_uri(data_uri_or_bytes)
    elif isinstance(data_uri_or_bytes, (bytes, bytearray)):
        raw_bytes = bytes(data_uri_or_bytes)
    else:
        return None

    svg_w, svg_h = _get_svg_dimensions(raw_bytes)
    package = doc.part.package

    from docx.parts.image import ImagePart
    from docx.opc.packuri import PackURI
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.shape import CT_Inline
    from docx.shared import Inches, Emu

    n = len(package.image_parts) + 1
    partname = PackURI(f"/word/media/image{n}.svg")
    image_part = ImagePart(partname, "image/svg+xml", raw_bytes)
    package.image_parts.append(image_part)

    rId = doc.part.relate_to(image_part, RT.IMAGE)
    shape_id = doc.part.next_id

    width_emu = Inches(width_inches).emu
    height_emu = int(width_emu * svg_h / svg_w)

    inline = CT_Inline.new_pic_inline(shape_id, rId, f"image{n}.svg", width_emu, height_emu)

    p = doc.add_paragraph()
    run = p.add_run()
    run._r.add_drawing(inline)
    return inline
